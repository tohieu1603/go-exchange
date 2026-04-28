# -*- coding: utf-8 -*-
"""Sinh bao cao do an tot nghiep .docx — Tieng Viet co dau, ~11 chuong.

Chay: python3 docs/report-assets/generate-report.py
Dau ra: docs/BaoCao-MicroExchange.docx
"""
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS = os.path.join(ROOT, "report-assets")
OUT = os.path.join(ROOT, "BaoCao-MicroExchange.docx")


# ═══════════════════════════════════════════════════════════════════
# TIEN ICH
# ═══════════════════════════════════════════════════════════════════
def _set_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name = name
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def P(doc, text, size=13, bold=False, italic=False,
      align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, after=4):
    """Doan van ban thuong."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    if indent:
        pf.first_line_indent = Cm(1.0)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    _set_font(r, size=size, bold=bold, italic=italic)
    return p


def H(doc, text, level=1):
    """Heading — level 0=title, 1=chapter, 2=section, 3=subsection."""
    sizes = {0: 18, 1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level <= 1 else 10)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    r = p.add_run(text)
    _set_font(r, size=sizes.get(level, 13), bold=True, color=(0x0B, 0x2A, 0x52))
    return p


def B(doc, text, size=13):
    """Bullet list item."""
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.text = ""
    r = p.add_run(text)
    _set_font(r, size=size)
    return p


def CAP(doc, text):
    """Caption cho hinh/bang."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    _set_font(r, size=12, italic=True, color=(0x33, 0x33, 0x33))
    return p


def IMG(doc, path, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))


def TBL(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(h)
        _set_font(r, size=11, bold=True, color=(0xFF, 0xFF, 0xFF))
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "0B2A52"); tcpr.append(shd)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = pp.add_run(str(val))
            _set_font(r, size=11)
    if col_widths:
        for row in tbl.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def PB(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# KHOI TAO
# ═══════════════════════════════════════════════════════════════════
doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(13)


# ═══════════════════════════════════════════════════════════════════
# TRANG BIA
# ═══════════════════════════════════════════════════════════════════
for _ in range(2): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO"); _set_font(r, size=14, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN"); _set_font(r, size=14, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("KHOA CÔNG NGHỆ PHẦN MỀM"); _set_font(r, size=13, bold=True)
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ĐỒ ÁN TỐT NGHIỆP"); _set_font(r, size=20, bold=True, color=(0x0B, 0x2A, 0x52))
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ĐỀ TÀI"); _set_font(r, size=14, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("XÂY DỰNG NỀN TẢNG GIAO DỊCH TIỀN MÃ HÓA\nTHEO KIẾN TRÚC VI DỊCH VỤ"); _set_font(r, size=18, bold=True, color=(0xA9, 0x1B, 0x1B))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("(MICRO-EXCHANGE)"); _set_font(r, size=16, bold=True, color=(0xA9, 0x1B, 0x1B))
for _ in range(3): doc.add_paragraph()
info = [
    ("Giảng viên hướng dẫn:", "..............................................."),
    ("Sinh viên thực hiện:", "..............................................."),
    ("Mã số sinh viên:", "..............................................."),
    ("Lớp:", "..............................................."),
    ("Niên khoá:", "2025 – 2026"),
]
for lb, val in info:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(lb + "  "); _set_font(r1, size=13, bold=True)
    r2 = p.add_run(val); _set_font(r2, size=13)
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Thành phố Hồ Chí Minh, tháng 04/2026"); _set_font(r, size=13, italic=True)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# LOI CAM DOAN
# ═══════════════════════════════════════════════════════════════════
H(doc, "LỜI CAM ĐOAN", 1)
P(doc, "Tôi xin cam đoan đồ án tốt nghiệp với đề tài \"Xây dựng nền tảng giao dịch tiền mã hoá theo kiến trúc vi dịch vụ (Micro-Exchange)\" là công trình nghiên cứu và thực hiện của cá nhân tôi dưới sự hướng dẫn của giảng viên. Các số liệu, kết quả trình bày trong đồ án là trung thực và chưa từng được công bố trong bất kỳ công trình nào khác.")
P(doc, "Tất cả các tài liệu tham khảo đã được trích dẫn đầy đủ. Tôi hoàn toàn chịu trách nhiệm về tính trung thực và nguyên bản của đồ án này.")
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("TP. Hồ Chí Minh, ngày ... tháng 04 năm 2026\nSinh viên thực hiện\n\n\n(Ký và ghi rõ họ tên)"); _set_font(r, size=13)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# LOI CAM ON
# ═══════════════════════════════════════════════════════════════════
H(doc, "LỜI CẢM ƠN", 1)
P(doc, "Lời đầu tiên, tôi xin gửi lời cảm ơn chân thành nhất đến giảng viên hướng dẫn đã tận tình chỉ bảo, định hướng và hỗ trợ tôi trong suốt quá trình thực hiện đồ án.")
P(doc, "Tôi cũng xin cảm ơn toàn thể quý thầy cô trong Khoa Công nghệ Phần mềm, Trường Đại học Công nghệ Thông tin đã truyền đạt những kiến thức nền tảng quý báu, giúp tôi có đủ năng lực để hoàn thành đề tài này.")
P(doc, "Cuối cùng, tôi xin cảm ơn gia đình và bạn bè đã luôn động viên, khích lệ trong suốt thời gian học tập và nghiên cứu.")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# NHAN XET GVHD / GVPB (de trong)
# ═══════════════════════════════════════════════════════════════════
H(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", 1)
for _ in range(25): P(doc, "...........................................................................................................................", indent=False, after=2)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Điểm đánh giá: ............\n\nGiảng viên hướng dẫn\n(Ký và ghi rõ họ tên)"); _set_font(r, size=13)
PB(doc)

H(doc, "NHẬN XÉT CỦA GIẢNG VIÊN PHẢN BIỆN", 1)
for _ in range(25): P(doc, "...........................................................................................................................", indent=False, after=2)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Điểm đánh giá: ............\n\nGiảng viên phản biện\n(Ký và ghi rõ họ tên)"); _set_font(r, size=13)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# MUC LUC
# ═══════════════════════════════════════════════════════════════════
H(doc, "MỤC LỤC", 1)
toc_items = [
    "LỜI CAM ĐOAN",
    "LỜI CẢM ƠN",
    "DANH MỤC HÌNH",
    "DANH MỤC BẢNG",
    "DANH MỤC TỪ VIẾT TẮT",
    "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI",
    "  1.1. Lý do chọn đề tài",
    "  1.2. Mục tiêu đề tài",
    "  1.3. Đối tượng và phạm vi nghiên cứu",
    "  1.4. Phương pháp nghiên cứu",
    "  1.5. Ý nghĩa khoa học và thực tiễn",
    "  1.6. Bố cục đồ án",
    "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT",
    "  2.1. Tổng quan về sàn giao dịch tiền mã hoá",
    "  2.2. Kiến trúc vi dịch vụ (Microservices)",
    "  2.3. Kiến trúc hướng sự kiện (EDA)",
    "  2.4. Mẫu CQRS",
    "  2.5. Database-per-Service",
    "  2.6. Order Book và thuật toán khớp lệnh",
    "  2.7. gRPC và Protocol Buffers",
    "  2.8. Redis: Cache, Lua Script, Streams, Pub/Sub",
    "  2.9. Apache Kafka",
    "  2.10. JWT, TOTP và bcrypt",
    "  2.11. WebSocket",
    "CHƯƠNG 3. KHẢO SÁT VÀ PHÂN TÍCH YÊU CẦU",
    "  3.1. Khảo sát các sàn giao dịch hiện có",
    "  3.2. Yêu cầu chức năng",
    "  3.3. Yêu cầu phi chức năng",
    "  3.4. Tác nhân hệ thống",
    "  3.5. Biểu đồ Use Case tổng hợp",
    "  3.6. Đặc tả chi tiết Use Case",
    "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG",
    "  4.1. Kiến trúc tổng thể",
    "  4.2. Danh mục vi dịch vụ",
    "  4.3. Thiết kế cơ sở dữ liệu",
    "  4.4. Thiết kế giao diện gRPC",
    "  4.5. Thiết kế API REST",
    "  4.6. Thiết kế Event Schema",
    "  4.7. Thiết kế kênh WebSocket",
    "  4.8. Thuật toán khớp lệnh",
    "  4.9. Thuật toán thanh lý vị thế Futures",
    "  4.10. Thiết kế Redis Lua Script cho số dư",
    "CHƯƠNG 5. NGĂN XẾP CÔNG NGHỆ",
    "  5.1. Bảng công nghệ chi tiết",
    "  5.2. Lý do lựa chọn",
    "  5.3. Cấu trúc thư mục dự án",
    "CHƯƠNG 6. HIỆN THỰC HỆ THỐNG",
    "  6.1. Auth Service",
    "  6.2. Wallet Service",
    "  6.3. Market Service",
    "  6.4. Trading Service và Matching Engine",
    "  6.5. Futures Service và Liquidation Engine",
    "  6.6. Notification Service",
    "  6.7. API Gateway",
    "  6.8. ES Indexer",
    "CHƯƠNG 7. LUỒNG NGHIỆP VỤ CHI TIẾT",
    "  7.1. Luồng đặt lệnh Spot",
    "  7.2. Luồng thanh lý vị thế Futures",
    "  7.3. Luồng nạp tiền VND qua SePay",
    "  7.4. Luồng rút tiền",
    "  7.5. Luồng xác thực và KYC",
    "CHƯƠNG 8. CƠ CHẾ EVENT-DRIVEN VÀ CQRS",
    "  8.1. Kiến trúc Event Bus",
    "  8.2. Mô hình CQRS trên hot path",
    "  8.3. Danh sách Domain Event",
    "CHƯƠNG 9. BẢO MẬT",
    "  9.1. Xác thực và phân quyền",
    "  9.2. Bảo vệ API",
    "  9.3. Chống gian lận",
    "  9.4. Ánh xạ OWASP Top 10",
    "CHƯƠNG 10. TRIỂN KHAI VÀ KHẢ NĂNG MỞ RỘNG",
    "  10.1. Sơ đồ triển khai Docker Compose",
    "  10.2. Biến môi trường",
    "  10.3. Khả năng mở rộng",
    "CHƯƠNG 11. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
    "  11.1. Kết quả đạt được",
    "  11.2. Hạn chế",
    "  11.3. Hướng phát triển",
    "PHỤ LỤC A. DANH MỤC TỪ VIẾT TẮT VÀ THUẬT NGỮ",
    "PHỤ LỤC B. MẪU JSON REQUEST/RESPONSE",
    "TÀI LIỆU THAM KHẢO",
]
for t in toc_items:
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(2)
    rr = pp.add_run(t); _set_font(rr, size=13)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# DANH MUC HINH
# ═══════════════════════════════════════════════════════════════════
H(doc, "DANH MỤC HÌNH", 1)
figs_list = [
    "Hình 1. Kiến trúc tổng thể nền tảng Micro-Exchange",
    "Hình 2. Luồng xử lý lệnh Spot theo mô hình CQRS",
    "Hình 3. Luồng mở và thanh lý vị thế Futures",
    "Hình 4. Sơ đồ Event Bus — Producer · Broker · Consumer",
    "Hình 5. Biểu đồ Use Case tổng hợp",
    "Hình 6. Sơ đồ tuần tự — Đặt lệnh giao dịch Spot",
    "Hình 7. Sơ đồ tuần tự — Thanh lý vị thế Futures",
    "Hình 8. Sơ đồ quan hệ thực thể (ERD) tổng hợp",
    "Hình 9. Sơ đồ triển khai Docker Compose",
    "Hình 10. Cấu trúc Order Book và thuật toán khớp lệnh",
]
for f in figs_list:
    P(doc, f, indent=False, after=2)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# DANH MUC BANG
# ═══════════════════════════════════════════════════════════════════
H(doc, "DANH MỤC BẢNG", 1)
tbls_list = [
    "Bảng 3.1. So sánh các sàn giao dịch tiền mã hoá",
    "Bảng 3.2. Danh sách yêu cầu chức năng",
    "Bảng 3.3. Yêu cầu phi chức năng",
    "Bảng 3.4. Danh sách Use Case",
    "Bảng 3.5. Đặc tả UC-07 — Đặt lệnh Spot",
    "Bảng 3.6. Đặc tả UC-09 — Mở vị thế Futures",
    "Bảng 4.1. Danh mục vi dịch vụ và trách nhiệm",
    "Bảng 4.2. Phân bổ CSDL theo vi dịch vụ",
    "Bảng 4.3. Giao diện gRPC",
    "Bảng 4.4. Danh sách API REST Endpoint",
    "Bảng 4.5. Danh sách Domain Event",
    "Bảng 4.6. Kênh WebSocket",
    "Bảng 5.1. Ngăn xếp công nghệ chi tiết",
    "Bảng 9.1. Ánh xạ OWASP Top 10",
    "Bảng 10.1. Biến môi trường chính",
    "Bảng A.1. Danh mục từ viết tắt và thuật ngữ",
]
for t in tbls_list:
    P(doc, t, indent=False, after=2)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# DANH MUC TU VIET TAT (ban ngan o dau)
# ═══════════════════════════════════════════════════════════════════
H(doc, "DANH MỤC TỪ VIẾT TẮT", 1)
TBL(doc,
    ["Từ viết tắt", "Đầy đủ"],
    [
        ["API",     "Application Programming Interface"],
        ["CQRS",    "Command Query Responsibility Segregation"],
        ["CSDL",    "Cơ sở dữ liệu"],
        ["EDA",     "Event-Driven Architecture"],
        ["ERD",     "Entity-Relationship Diagram"],
        ["gRPC",    "gRPC Remote Procedure Call"],
        ["JWT",     "JSON Web Token"],
        ["KYC",     "Know Your Customer"],
        ["ORM",     "Object-Relational Mapping"],
        ["PnL",     "Profit and Loss"],
        ["REST",    "Representational State Transfer"],
        ["TOTP",    "Time-based One-Time Password"],
        ["WS",      "WebSocket"],
    ],
    col_widths=[4.0, 12.0],
)
P(doc, "Danh sách đầy đủ xem tại Phụ lục A.", italic=True)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 1. TONG QUAN DE TAI
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)

H(doc, "1.1. Lý do chọn đề tài", 2)
P(doc, "Thị trường tiền mã hoá (cryptocurrency) đã và đang phát triển mạnh mẽ trên phạm vi toàn cầu. Theo thống kê từ CoinMarketCap (2025), tổng vốn hoá thị trường crypto vượt ngưỡng 3 nghìn tỷ USD với hơn 500 triệu người dùng. Các sàn giao dịch (exchange) lớn như Binance, Coinbase, OKX và Bybit đều vận hành trên nền tảng phân tán phức tạp, phục vụ hàng triệu người dùng đồng thời với yêu cầu khắt khe về độ trễ xử lý lệnh — thường ở mức vài mili giây.")
P(doc, "Để đáp ứng được khối lượng giao dịch lớn và yêu cầu real-time, các hệ thống thực tế thường áp dụng kiến trúc vi dịch vụ (microservices) kết hợp với mô hình hướng sự kiện (Event-Driven Architecture — EDA) và tách riêng đường xử lý ghi (command) với đường đọc (query) theo nguyên tắc CQRS (Command Query Responsibility Segregation).")
P(doc, "Từ nhu cầu thực tế đó, đề tài \"Micro-Exchange\" được xây dựng nhằm mục đích học thuật: tái hiện lại kiến trúc của một sàn giao dịch thực trên quy mô nhỏ, qua đó làm rõ các khái niệm về microservices, giao tiếp qua gRPC, hàng đợi sự kiện, order book khớp lệnh trong bộ nhớ (in-memory) và cơ chế quản lý số dư an toàn bằng Redis Lua script — tất cả được đóng gói trong một hệ thống có thể chạy trên môi trường Docker Compose cục bộ.")
P(doc, "Đề tài không chỉ giải quyết bài toán kỹ thuật mà còn là cơ hội để sinh viên vận dụng tổng hợp kiến thức từ nhiều môn học: Kiến trúc phần mềm, Cơ sở dữ liệu phân tán, Mạng máy tính, An toàn thông tin và Lập trình hệ thống.")

H(doc, "1.2. Mục tiêu đề tài", 2)
P(doc, "Đề tài hướng tới các mục tiêu cụ thể sau:")
B(doc, "Xây dựng hệ thống giao dịch giao ngay (Spot) cho phép người dùng đặt lệnh LIMIT và MARKET trên các cặp giao dịch ghép với USDT, hỗ trợ 55 cặp giao dịch bao gồm cả crypto, forex và commodity.")
B(doc, "Hỗ trợ sản phẩm hợp đồng tương lai vĩnh viễn (Perpetual Futures) với đòn bẩy từ 1x đến 125x, có cơ chế thanh lý (liquidation) tự động chạy mỗi giây.")
B(doc, "Áp dụng nguyên tắc database-per-service — mỗi vi dịch vụ sở hữu riêng một CSDL PostgreSQL, giao tiếp nội bộ qua gRPC và truyền sự kiện qua Redis Streams kết hợp Apache Kafka.")
B(doc, "Đảm bảo hot path xử lý lệnh không chạm CSDL quan hệ (zero DB write on critical path), chỉ thao tác trên in-memory order book và Redis Lua script, sau đó phát sự kiện để projector bất đồng bộ ghi vào DB.")
B(doc, "Cung cấp dữ liệu realtime (giá, độ sâu sổ lệnh, thông báo) tới client qua WebSocket, nhận giá thị trường từ Bybit WebSocket.")
B(doc, "Hiện thực đầy đủ các chức năng hỗ trợ: xác thực JWT, 2FA TOTP, KYC 4 bước, nạp/rút tiền VND (tích hợp SePay QR), quản lý ví đa tiền tệ, thông báo, giám sát gian lận và bảng điều khiển quản trị.")

H(doc, "1.3. Đối tượng và phạm vi nghiên cứu", 2)
P(doc, "Đối tượng nghiên cứu: Kiến trúc và kỹ thuật xây dựng nền tảng giao dịch tiền mã hoá theo mô hình vi dịch vụ, bao gồm cơ chế khớp lệnh, quản lý vị thế phái sinh, đồng bộ số dư qua event bus và bảo mật giao dịch.")
P(doc, "Phạm vi:")
B(doc, "Backend: toàn bộ 8 vi dịch vụ (auth, wallet, market, trading, futures, notification, gateway, es-indexer) cùng thư viện dùng chung (shared).")
B(doc, "Infrastructure: Redis, Kafka, Elasticsearch, PostgreSQL — quản lý qua Docker Compose.")
B(doc, "Frontend: không nằm trong phạm vi báo cáo; tuy nhiên hệ thống cung cấp đầy đủ API REST và kênh WebSocket để frontend tích hợp.")
B(doc, "Thanh toán: tích hợp SePay (VND) ở mức cơ bản; các quy trình pháp lý KYC/AML được mô phỏng bằng dữ liệu thử nghiệm.")

H(doc, "1.4. Phương pháp nghiên cứu", 2)
B(doc, "Nghiên cứu lý thuyết: Tham khảo sách và tài liệu chuyên ngành về microservices (Newman, 2021), CQRS (Fowler), thiết kế hệ thống dữ liệu chuyên sâu (Kleppmann, 2017).")
B(doc, "Khảo sát thực tế: Phân tích kiến trúc và API công khai của Binance, Coinbase, OKX, Bybit để rút ra các mẫu thiết kế phù hợp.")
B(doc, "Thực nghiệm: Xây dựng hệ thống hoàn chỉnh bằng Go, chạy thử với dữ liệu giá thực từ Bybit WebSocket, đánh giá hiệu năng trên máy phát triển.")

H(doc, "1.5. Ý nghĩa khoa học và thực tiễn", 2)
P(doc, "Về mặt khoa học, đồ án minh hoạ cách áp dụng tổng hợp các mẫu kiến trúc phân tán hiện đại (microservices, CQRS, EDA, database-per-service) vào một bài toán cụ thể có tính phức tạp cao — sàn giao dịch tài chính. Về mặt thực tiễn, sản phẩm có thể dùng làm nền tảng tham khảo (reference implementation) cho các đội phát triển muốn xây dựng hệ thống tương tự, hoặc làm tài liệu giảng dạy trong các môn học Kiến trúc phần mềm phân tán.")

H(doc, "1.6. Bố cục đồ án", 2)
P(doc, "Đồ án được tổ chức thành 11 chương:")
B(doc, "Chương 1 — Tổng quan: giới thiệu lý do, mục tiêu, phạm vi và phương pháp.")
B(doc, "Chương 2 — Cơ sở lý thuyết: trình bày nền tảng lý thuyết về microservices, CQRS, EDA, gRPC, Redis, Kafka, JWT, TOTP, WebSocket.")
B(doc, "Chương 3 — Phân tích yêu cầu: khảo sát, danh sách use case, đặc tả chi tiết.")
B(doc, "Chương 4 — Thiết kế hệ thống: kiến trúc, ERD, API, gRPC, thuật toán khớp lệnh.")
B(doc, "Chương 5 — Công nghệ: bảng tech stack, lý do lựa chọn, cấu trúc thư mục.")
B(doc, "Chương 6 — Hiện thực: chi tiết triển khai từng vi dịch vụ.")
B(doc, "Chương 7 — Luồng nghiệp vụ: mô tả chi tiết các luồng chính kèm sơ đồ tuần tự.")
B(doc, "Chương 8 — Event-Driven và CQRS: event bus, domain event, projector.")
B(doc, "Chương 9 — Bảo mật: xác thực, phân quyền, chống gian lận, OWASP Top 10.")
B(doc, "Chương 10 — Triển khai: Docker Compose, biến môi trường, khả năng mở rộng.")
B(doc, "Chương 11 — Kết luận: kết quả, hạn chế và hướng phát triển.")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 2. CO SO LY THUYET
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", 1)

H(doc, "2.1. Tổng quan về sàn giao dịch tiền mã hoá", 2)
P(doc, "Sàn giao dịch tiền mã hoá (cryptocurrency exchange) là nền tảng trung gian cho phép người dùng mua, bán và giao dịch các tài sản số (digital asset) như Bitcoin, Ethereum, Solana, v.v. Có hai loại sàn chính:")
B(doc, "Sàn tập trung (Centralized Exchange — CEX): vận hành bởi một tổ chức duy nhất, lưu giữ tài sản của người dùng (custodial), có sổ lệnh (order book) tập trung. Ví dụ: Binance, Coinbase, OKX, Bybit.")
B(doc, "Sàn phi tập trung (Decentralized Exchange — DEX): hoạt động trên blockchain, sử dụng hợp đồng thông minh (smart contract), không giữ hộ tài sản. Ví dụ: Uniswap, PancakeSwap.")
P(doc, "Đề tài Micro-Exchange thuộc loại CEX, tập trung vào hai sản phẩm giao dịch chính: giao dịch giao ngay (Spot) và hợp đồng tương lai vĩnh viễn (Perpetual Futures).")
P(doc, "Giao dịch Spot là hình thức mua bán thực tại giá hiện hành. Khi lệnh được khớp, tài sản được chuyển ngay lập tức giữa người mua và người bán. Hai loại lệnh phổ biến: lệnh giới hạn (LIMIT) — chỉ khớp tại giá do người dùng chỉ định; và lệnh thị trường (MARKET) — khớp ngay tại giá tốt nhất hiện có trên sổ lệnh.")
P(doc, "Perpetual Futures là hợp đồng phái sinh cho phép đặt cược về hướng giá mà không sở hữu tài sản cơ sở, không có ngày đáo hạn. Người dùng ký quỹ (margin) một phần giá trị hợp đồng và sử dụng đòn bẩy (leverage) để phóng đại lợi nhuận/thua lỗ. Khi thua lỗ vượt ngưỡng ký quỹ, vị thế bị thanh lý (liquidation) tự động.")

H(doc, "2.2. Kiến trúc vi dịch vụ (Microservices)", 2)
P(doc, "Kiến trúc vi dịch vụ (microservices architecture) là phong cách thiết kế phần mềm trong đó một ứng dụng được phân rã thành tập hợp các dịch vụ nhỏ, độc lập, mỗi dịch vụ chạy trong tiến trình riêng và giao tiếp qua cơ chế nhẹ (lightweight mechanism) như HTTP/REST hoặc gRPC (Newman, 2021). Các đặc trưng chính bao gồm:")
B(doc, "Triển khai độc lập (independent deployment): mỗi dịch vụ có thể build, test và deploy riêng mà không ảnh hưởng đến dịch vụ khác.")
B(doc, "Sở hữu dữ liệu riêng (data ownership): mỗi dịch vụ quản lý CSDL riêng, không chia sẻ bảng với dịch vụ khác — đây là nguyên tắc database-per-service.")
B(doc, "Giao tiếp qua API: giao tiếp đồng bộ qua gRPC/REST, bất đồng bộ qua message broker (Kafka, Redis Streams).")
B(doc, "Khả năng mở rộng theo chiều ngang (horizontal scalability): mỗi dịch vụ có thể scale độc lập theo nhu cầu riêng.")
P(doc, "So với kiến trúc nguyên khối (monolith), microservices có ưu điểm về tính mô-đun, khả năng mở rộng và tốc độ phát triển nhóm. Tuy nhiên, nó cũng mang lại độ phức tạp cao hơn về vận hành, giám sát và đảm bảo tính nhất quán dữ liệu phân tán.")

H(doc, "2.3. Kiến trúc hướng sự kiện (Event-Driven Architecture)", 2)
P(doc, "Kiến trúc hướng sự kiện (Event-Driven Architecture — EDA) là mô hình trong đó các thành phần giao tiếp thông qua việc phát (publish) và tiêu thụ (consume) sự kiện (event). Một sự kiện là một thông điệp bất biến (immutable message) mô tả một sự thật đã xảy ra trong hệ thống — ví dụ: \"một giao dịch đã được thực thi\" (trade.executed), \"số dư đã thay đổi\" (balance.changed).")
P(doc, "EDA mang lại nhiều lợi ích cho hệ thống phân tán:")
B(doc, "Giảm ghép nối (loose coupling): producer không cần biết consumer là ai, chỉ cần phát sự kiện lên bus.")
B(doc, "Khả năng mở rộng: thêm consumer mới mà không cần sửa producer.")
B(doc, "Khả năng phục hồi: sự kiện được lưu trữ bền vững trong Kafka/Redis Streams, consumer có thể replay từ bất kỳ offset nào.")
B(doc, "Phân tách trách nhiệm: mỗi consumer xử lý logic nghiệp vụ riêng — projector ghi DB, es-indexer đánh chỉ mục, notification gửi thông báo.")

H(doc, "2.4. Mẫu CQRS (Command Query Responsibility Segregation)", 2)
P(doc, "CQRS là mẫu kiến trúc tách biệt hoàn toàn đường xử lý ghi (command) và đường đọc (query) thành hai pipeline khác nhau, mỗi pipeline có thể tối ưu độc lập (Fowler). Trong ngữ cảnh sàn giao dịch:")
B(doc, "Command path (hot path): xử lý lệnh đặt/huỷ trên in-memory order book, cập nhật số dư qua Redis Lua script (atomic), sau đó phát domain event. Không ghi CSDL quan hệ trên đường này.")
B(doc, "Query path: phục vụ truy vấn lịch sử giao dịch, lịch sử lệnh, báo cáo — đọc từ PostgreSQL và Elasticsearch, được cập nhật bởi projector tiêu thụ sự kiện.")
P(doc, "Ưu điểm: hot path có độ trễ cực thấp (sub-millisecond) vì không chờ disk I/O; read model có thể tối ưu chỉ mục, phân trang và full-text search riêng biệt.")

H(doc, "2.5. Database-per-Service", 2)
P(doc, "Nguyên tắc database-per-service quy định mỗi vi dịch vụ sở hữu riêng một CSDL, không truy cập trực tiếp CSDL của dịch vụ khác. Mọi nhu cầu dữ liệu chéo phải qua API (gRPC — đồng bộ) hoặc qua sự kiện (bất đồng bộ). Trong Micro-Exchange, hệ thống sử dụng 6 instance PostgreSQL 16 riêng biệt trên các cổng 5551–5556.")
P(doc, "Ưu điểm: giảm coupling, cho phép mỗi dịch vụ tự chọn schema, migration, thời điểm nâng cấp. Nhược điểm: không có khoá ngoại vật lý liên-CSDL, phải đảm bảo nhất quán cuối cùng (eventual consistency) qua sự kiện.")

H(doc, "2.6. Order Book và thuật toán khớp lệnh", 2)
P(doc, "Order Book (sổ lệnh) là cấu trúc dữ liệu trung tâm của mọi sàn giao dịch, lưu trữ tất cả các lệnh mua (bid) và bán (ask) đang chờ khớp cho một cặp giao dịch. Mỗi bên được tổ chức thành các mức giá (price level), mỗi mức giá chứa danh sách lệnh theo thứ tự thời gian (FIFO).")
P(doc, "Thuật toán khớp lệnh phổ biến nhất là Price-Time Priority:")
B(doc, "Bước 1: Khi lệnh mới (incoming) đến, so sánh với mức giá tốt nhất của bên đối (best ask nếu lệnh mua, best bid nếu lệnh bán).")
B(doc, "Bước 2: Nếu giá thoả mãn (lệnh MARKET luôn thoả, lệnh LIMIT chỉ thoả khi giá đối phương nằm trong phạm vi), tiến hành khớp.")
B(doc, "Bước 3: Khớp theo thứ tự FIFO trong cùng mức giá. Số lượng khớp = min(còn lại của incoming, còn lại của resting).")
B(doc, "Bước 4: Cập nhật filledAmount cả hai phía. Nếu lệnh resting đã đầy (filled), loại khỏi sổ lệnh.")
B(doc, "Bước 5: Lặp lại cho đến khi incoming hết số lượng hoặc không còn giá thoả mãn.")
P(doc, "Trong Micro-Exchange, order book được triển khai hoàn toàn trong bộ nhớ (in-memory) bằng Go slice, sắp xếp bids giảm dần và asks tăng dần. Dữ liệu chỉ được phục hồi từ DB khi service khởi động lại (cold-start recovery).")

H(doc, "2.7. gRPC và Protocol Buffers", 2)
P(doc, "gRPC là framework RPC (Remote Procedure Call) hiệu năng cao do Google phát triển, sử dụng HTTP/2 làm giao thức truyền tải và Protocol Buffers (protobuf) làm định dạng tuần tự hoá. So với REST/JSON, gRPC mang lại:")
B(doc, "Tốc độ tuần tự hoá/giải tuần tự hoá nhanh hơn 5–10 lần nhờ định dạng nhị phân.")
B(doc, "Contract-first: giao diện dịch vụ được định nghĩa trong file .proto, compiler tự sinh code cho cả client và server.")
B(doc, "HTTP/2: hỗ trợ multiplexing, header compression, bidirectional streaming.")
P(doc, "Trong Micro-Exchange, gRPC được sử dụng cho giao tiếp nội bộ giữa các dịch vụ: Trading → Wallet (LockBalance, CheckBalance), Futures → Wallet (Deduct, Credit), Gateway → Auth (ValidateToken), Trading → Market (GetPrice).")

H(doc, "2.8. Redis: Cache, Lua Script, Streams, Pub/Sub", 2)
P(doc, "Redis là hệ thống lưu trữ dữ liệu in-memory, đa mục đích. Trong Micro-Exchange, Redis đóng 4 vai trò đồng thời:")
B(doc, "Cache số dư: Lưu trữ số dư (bal:userID:CUR) và số dư khoá (locked:userID:CUR) dưới dạng key-value, độ trễ đọc/ghi sub-millisecond.")
B(doc, "Lua Script atomic: Các thao tác kiểm tra-rồi-trừ (check-then-deduct) và kiểm tra-rồi-khoá (check-then-lock) được đóng gói trong Lua script, đảm bảo tính nguyên tử mà không cần lock phân tán.")
B(doc, "Streams (event bus hot-path): Redis Streams hoạt động như message queue nhẹ, mỗi sự kiện được ghi dưới dạng entry trong stream, hỗ trợ consumer group để load balance.")
B(doc, "Pub/Sub: Dùng cho WebSocket fan-out — khi có giao dịch mới, Trading Service publish lên kênh ws:broadcast, Gateway Hub subscribe và đẩy tới tất cả client đang kết nối.")

H(doc, "2.9. Apache Kafka", 2)
P(doc, "Apache Kafka là nền tảng streaming phân tán, hoạt động như backbone sự kiện bền vững (durable event backbone). Kafka lưu trữ sự kiện theo topic, mỗi topic chia thành nhiều partition cho phép xử lý song song. Consumer group đảm bảo mỗi sự kiện chỉ được xử lý một lần trong cùng nhóm (at-least-once), và nhiều nhóm khác nhau đều nhận bản sao (fan-out).")
P(doc, "Trong Micro-Exchange, Kafka 3.9 chạy ở chế độ KRaft (không cần ZooKeeper) với một broker duy nhất (phù hợp môi trường phát triển). Các topic chính: trade.executed, order.updated, balance.changed, position.changed, user.registered.")

H(doc, "2.10. JWT, TOTP và bcrypt", 2)
P(doc, "JWT (JSON Web Token) là chuẩn mở RFC 7519 cho phép truyền tải thông tin xác thực dưới dạng token có chữ ký số. Trong Micro-Exchange, Auth Service phát hành JWT ký bằng HS256, bao gồm: userID, email, role, thời gian hết hạn. Gateway xác thực JWT qua gRPC call tới Auth Service trước khi cho phép request đi tiếp.")
P(doc, "TOTP (Time-based One-Time Password, RFC 6238) là cơ chế mật khẩu dùng một lần dựa trên thời gian, tạo mã 6 chữ số mới mỗi 30 giây. Người dùng quét mã QR bằng Google Authenticator hoặc Authy để kích hoạt 2FA. Secret TOTP được lưu mã hoá trong DB, không bao giờ lộ ra ngoài API (json tag: \"-\").")
P(doc, "bcrypt là hàm băm mật khẩu chậm có chủ đích (cost factor = 10), chống brute-force hiệu quả. Mật khẩu người dùng luôn được băm trước khi lưu, không bao giờ lưu dạng cleartext.")

H(doc, "2.11. WebSocket", 2)
P(doc, "WebSocket (RFC 6455) là giao thức truyền thông hai chiều (full-duplex) trên một kết nối TCP duy nhất. So với HTTP polling, WebSocket giảm đáng kể overhead vì không cần thiết lập kết nối mới cho mỗi lần cập nhật. Trong Micro-Exchange, WebSocket được dùng để đẩy real-time:")
B(doc, "Giá và biến động 24h (tickers@pair)")
B(doc, "Độ sâu sổ lệnh (depth@pair)")
B(doc, "Giao dịch mới (trades@pair)")
B(doc, "Biểu đồ nến (candles@pair@interval)")
B(doc, "Thông báo cá nhân (notifications@userID)")
B(doc, "Vị thế futures (positions@userID)")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 3. PHAN TICH YEU CAU
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 3. KHẢO SÁT VÀ PHÂN TÍCH YÊU CẦU", 1)

H(doc, "3.1. Khảo sát các sàn giao dịch hiện có", 2)
P(doc, "Để rút ra các mẫu thiết kế phù hợp, nhóm đã khảo sát 4 sàn giao dịch hàng đầu thế giới:")
TBL(doc,
    ["Tiêu chí", "Binance", "Coinbase", "OKX", "Bybit"],
    [
        ["Kiến trúc", "Microservices", "Microservices", "Microservices", "Microservices"],
        ["Order Book", "In-memory", "In-memory", "In-memory", "In-memory"],
        ["Event Bus", "Kafka + RocketMQ", "Kafka", "Kafka", "Kafka"],
        ["Số cặp Spot", ">2.000", ">500", ">700", ">1.200"],
        ["Futures", "Có (125x)", "Không", "Có (100x)", "Có (100x)"],
        ["Matching latency", "<1ms", "<5ms", "<1ms", "<1ms"],
        ["API", "REST + WS", "REST + WS", "REST + WS", "REST + WS"],
        ["Ngôn ngữ", "Java, Go, C++", "Go, Ruby", "Go, Java", "Go, C++"],
    ],
    col_widths=[3.5, 3.0, 3.0, 3.0, 3.0],
)
CAP(doc, "Bảng 3.1. So sánh các sàn giao dịch tiền mã hoá hàng đầu")
P(doc, "Nhận xét chung: tất cả các sàn đều sử dụng kiến trúc microservices, order book trong bộ nhớ và event bus (chủ yếu Kafka). Go là ngôn ngữ phổ biến nhờ hiệu năng cao và hỗ trợ concurrency tốt. Micro-Exchange kế thừa các mẫu này ở quy mô thu nhỏ phù hợp học thuật.")

H(doc, "3.2. Yêu cầu chức năng", 2)
TBL(doc,
    ["STT", "Mã", "Yêu cầu", "Mức ưu tiên"],
    [
        ["1", "FR-01", "Đăng ký tài khoản bằng email, xác thực email", "Cao"],
        ["2", "FR-02", "Đăng nhập, phát hành JWT, refresh token", "Cao"],
        ["3", "FR-03", "Bật/tắt xác thực hai yếu tố (2FA TOTP)", "Cao"],
        ["4", "FR-04", "KYC 4 bước: xác thực email → hồ sơ → giấy tờ → duyệt", "Trung bình"],
        ["5", "FR-05", "Xem số dư ví đa tiền tệ (VND, USDT, BTC, ETH, ...)", "Cao"],
        ["6", "FR-06", "Nạp VND qua SePay QR chuyển khoản ngân hàng", "Cao"],
        ["7", "FR-07", "Rút tiền VND về ngân hàng (Admin duyệt)", "Cao"],
        ["8", "FR-08", "Xem bảng giá realtime, biểu đồ nến (candlestick)", "Cao"],
        ["9", "FR-09", "Đặt lệnh Spot: LIMIT, MARKET", "Cao"],
        ["10", "FR-10", "Huỷ lệnh Spot chờ khớp", "Cao"],
        ["11", "FR-11", "Xem lịch sử lệnh và giao dịch", "Trung bình"],
        ["12", "FR-12", "Mở vị thế Futures với đòn bẩy 1–125x", "Cao"],
        ["13", "FR-13", "Đóng vị thế Futures thủ công", "Cao"],
        ["14", "FR-14", "Thiết lập Take Profit / Stop Loss", "Trung bình"],
        ["15", "FR-15", "Thanh lý (liquidation) tự động khi PnL vượt ngưỡng", "Cao"],
        ["16", "FR-16", "Nhận thông báo realtime (lệnh khớp, thanh lý, nạp tiền)", "Trung bình"],
        ["17", "FR-17", "Admin: quản lý người dùng, duyệt KYC, duyệt rút tiền", "Cao"],
        ["18", "FR-18", "Admin: cấu hình phí, quản lý cặp giao dịch, bonus", "Trung bình"],
        ["19", "FR-19", "Admin: giám sát nhật ký gian lận, khoá tài khoản", "Trung bình"],
    ],
    col_widths=[1.0, 1.5, 9.5, 3.0],
)
CAP(doc, "Bảng 3.2. Danh sách yêu cầu chức năng")

H(doc, "3.3. Yêu cầu phi chức năng", 2)
TBL(doc,
    ["STT", "Mã", "Yêu cầu", "Tiêu chí đo lường"],
    [
        ["1", "NFR-01", "Hiệu năng khớp lệnh", "Matching engine xử lý ≥10.000 lệnh/giây trên một cặp"],
        ["2", "NFR-02", "Độ trễ hot path", "< 5ms từ khi nhận lệnh đến khi phát sự kiện"],
        ["3", "NFR-03", "Tính sẵn sàng", "Mỗi service có thể restart độc lập, cold-start < 3s"],
        ["4", "NFR-04", "Tính nhất quán", "Eventual consistency qua event bus, DB sync < 1s"],
        ["5", "NFR-05", "Bảo mật", "Mật khẩu bcrypt, JWT HS256, 2FA TOTP, rate limit"],
        ["6", "NFR-06", "Khả năng mở rộng", "Horizontal scale per-service, per-pair sharding"],
        ["7", "NFR-07", "Khả năng giám sát", "Structured logging, health endpoint, consumer lag check"],
        ["8", "NFR-08", "Triển khai", "Docker Compose single command, hỗ trợ CI/CD"],
    ],
    col_widths=[1.0, 1.8, 5.0, 7.5],
)
CAP(doc, "Bảng 3.3. Yêu cầu phi chức năng")

H(doc, "3.4. Tác nhân hệ thống", 2)
P(doc, "Hệ thống xác định hai nhóm tác nhân (actor) chính:")
B(doc, "Người dùng cuối (End User): người dùng thông thường có nhu cầu nạp tiền, giao dịch Spot/Futures và rút tiền về tài khoản ngân hàng. Yêu cầu phải đăng ký, đăng nhập, và tuỳ theo mức KYC mà được phép sử dụng các tính năng nâng cao.")
B(doc, "Quản trị viên (Admin): nhân viên vận hành sàn, chịu trách nhiệm duyệt hồ sơ KYC, cấu hình phí giao dịch, duyệt yêu cầu rút tiền, giám sát gian lận và quản lý hệ thống.")
P(doc, "Ngoài ra, hệ thống tương tác với các tác nhân bên ngoài (external actor):")
B(doc, "Bybit WebSocket API: cung cấp giá thị trường real-time cho Market Service.")
B(doc, "SePay Payment Gateway: xử lý webhook xác nhận thanh toán nạp VND.")

H(doc, "3.5. Biểu đồ Use Case tổng hợp", 2)
P(doc, "Hình 5 trình bày biểu đồ use case tổng hợp, thể hiện ranh giới hệ thống và quan hệ giữa tác nhân với các ca sử dụng. Phần trên là các use case dành cho người dùng cuối (12 UC), phần dưới là use case dành cho quản trị viên (4 UC).")
IMG(doc, os.path.join(ASSETS, "fig5-usecase.png"), 15.5)
CAP(doc, "Hình 5. Biểu đồ Use Case tổng hợp của hệ thống Micro-Exchange")

TBL(doc,
    ["STT", "Mã UC", "Tên ca sử dụng", "Tác nhân"],
    [
        ["1",  "UC-01", "Đăng ký tài khoản",                    "End User"],
        ["2",  "UC-02", "Đăng nhập + xác thực 2FA (TOTP)",      "End User"],
        ["3",  "UC-03", "Nộp hồ sơ KYC",                        "End User"],
        ["4",  "UC-04", "Nạp tiền VND qua SePay QR",            "End User"],
        ["5",  "UC-05", "Nạp USDT on-chain",                    "End User"],
        ["6",  "UC-06", "Xem bảng giá và biểu đồ nến",          "End User"],
        ["7",  "UC-07", "Đặt lệnh Spot (LIMIT / MARKET)",       "End User"],
        ["8",  "UC-08", "Huỷ lệnh chờ khớp",                    "End User"],
        ["9",  "UC-09", "Mở vị thế Futures có đòn bẩy",         "End User"],
        ["10", "UC-10", "Đóng vị thế Futures thủ công",          "End User"],
        ["11", "UC-11", "Rút tiền về ngân hàng",                 "End User"],
        ["12", "UC-12", "Nhận thông báo realtime",               "End User"],
        ["13", "UC-13", "Duyệt hồ sơ KYC",                      "Admin"],
        ["14", "UC-14", "Cấu hình phí giao dịch, cặp coin",     "Admin"],
        ["15", "UC-15", "Duyệt yêu cầu rút tiền",               "Admin"],
        ["16", "UC-16", "Giám sát nhật ký gian lận",             "Admin"],
    ],
    col_widths=[1.0, 1.8, 7.5, 3.5],
)
CAP(doc, "Bảng 3.4. Danh sách Use Case")

H(doc, "3.6. Đặc tả chi tiết Use Case", 2)
P(doc, "Dưới đây là đặc tả chi tiết hai ca sử dụng trọng tâm nhất — đặt lệnh Spot và mở vị thế Futures.")

TBL(doc,
    ["Mục", "Nội dung"],
    [
        ["Mã",          "UC-07"],
        ["Tên",         "Đặt lệnh giao dịch giao ngay (Spot)"],
        ["Tác nhân",    "Người dùng cuối (đã đăng nhập, xác thực 2FA nếu bật)"],
        ["Mô tả",       "Cho phép người dùng mua/bán tiền mã hoá theo cặp (VD: BTC_USDT) với lệnh giới hạn (LIMIT) hoặc lệnh thị trường (MARKET)."],
        ["Tiền điều kiện", "1. Tài khoản đã kích hoạt.\n2. Có đủ số dư khả dụng cho cặp tiền tương ứng.\n3. Cặp giao dịch đang mở (isActive = true)."],
        ["Hậu điều kiện",  "1. Lệnh được ghi nhận trong hệ thống.\n2. Số dư được khoá tương ứng (LIMIT).\n3. Sự kiện trade.executed / order.updated được phát ra.\n4. Client nhận được cập nhật realtime qua WebSocket."],
        ["Luồng chính",
         "1. Người dùng chọn cặp, nhập loại lệnh (LIMIT/MARKET), phía (BUY/SELL), số lượng và giá (nếu LIMIT).\n"
         "2. Client gửi POST /api/trading/orders tới Gateway.\n"
         "3. Gateway xác thực JWT qua gRPC ValidateToken → Auth Service.\n"
         "4. Gateway chuyển tiếp request tới Trading Handler.\n"
         "5. Trading Handler validate payload (side, type, pair, amount > 0).\n"
         "6. Nếu LIMIT/BUY: gọi gRPC LockBalance → Wallet Service để khoá USDT.\n"
         "   Nếu LIMIT/SELL: gọi gRPC LockBalance → Wallet Service để khoá base coin.\n"
         "7. Tạo đối tượng Order, ghi DB (pg-trading) với status = OPEN.\n"
         "8. Gọi MatchingEngine.ProcessOrder(order).\n"
         "9. OrderBook.Match() khớp lệnh trong bộ nhớ theo Price-Time Priority.\n"
         "10. Cho mỗi trade khớp: Redis Lua BalanceDeduct/BalanceCredit (atomic).\n"
         "11. Unlock số dư đã khoá cho phần đã khớp.\n"
         "12. Phát sự kiện: trade.executed, order.updated, balance.changed qua Redis Streams.\n"
         "13. WebSocket Hub broadcast trades@pair, depth@pair.\n"
         "14. Trả về HTTP 200 với orderId, status và danh sách trades."],
        ["Luồng phụ",
         "5a. Nếu validate thất bại: trả HTTP 400 với thông báo lỗi.\n"
         "6a. Nếu số dư không đủ: trả HTTP 400 \"insufficient balance\".\n"
         "9a. Nếu LIMIT không tìm được đối khớp: thêm vào sổ lệnh chờ, trả status = OPEN.\n"
         "10a. Nếu Redis lỗi: rollback order book, trả HTTP 500."],
    ],
    col_widths=[3.0, 13.0],
)
CAP(doc, "Bảng 3.5. Đặc tả Use Case UC-07 — Đặt lệnh Spot")

TBL(doc,
    ["Mục", "Nội dung"],
    [
        ["Mã",          "UC-09"],
        ["Tên",         "Mở vị thế Futures có đòn bẩy"],
        ["Tác nhân",    "Người dùng cuối"],
        ["Mô tả",       "Cho phép người dùng tạo vị thế long/short trên hợp đồng tương lai vĩnh viễn, với đòn bẩy từ 1x đến 125x."],
        ["Tiền điều kiện",
         "1. Đăng nhập thành công.\n2. Có số dư USDT đủ cho margin + phí.\n3. Cặp futures hợp lệ."],
        ["Hậu điều kiện",
         "1. Vị thế được tạo ở trạng thái OPEN.\n2. Margin bị trừ khỏi số dư.\n3. Liquidation Engine bắt đầu theo dõi giá mark.\n4. Sự kiện position.changed được phát."],
        ["Luồng chính",
         "1. Người dùng gửi POST /api/futures/order {pair, side, size, leverage, takeProfit, stopLoss}.\n"
         "2. Futures Service truy giá mark từ Redis (price:BTC_USDT).\n"
         "3. Tính margin = (size × markPrice) / leverage.\n"
         "4. Tính phí = size × markPrice × 0.05%.\n"
         "5. Tính giá thanh lý: LONG → entry × (1 − 1/lev + 0.005), SHORT → entry × (1 + 1/lev − 0.005).\n"
         "6. Gọi gRPC CheckBalance(userID, USDT, margin + fee) → Wallet Service.\n"
         "7. Gọi gRPC Deduct(userID, USDT, margin + fee).\n"
         "8. Ghi record vào pg-futures {status: OPEN, entry_price, liquidation_price, ...}.\n"
         "9. Phát sự kiện position.changed.\n"
         "10. Gửi notification và broadcast qua WebSocket."],
        ["Luồng phụ",
         "2a. Giá không khả dụng → trả HTTP 400 \"price unavailable\".\n"
         "6a. Số dư không đủ → trả HTTP 400 \"insufficient balance\".\n"
         "8a. Lỗi DB → rollback gRPC Credit (hoàn trả margin)."],
    ],
    col_widths=[3.0, 13.0],
)
CAP(doc, "Bảng 3.6. Đặc tả Use Case UC-09 — Mở vị thế Futures")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 4. THIET KE HE THONG
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG", 1)

H(doc, "4.1. Kiến trúc tổng thể", 2)
P(doc, "Hệ thống Micro-Exchange được thiết kế theo kiến trúc phân lớp (layered microservices architecture) bao gồm 5 tầng chính, như minh hoạ tại Hình 1:")
B(doc, "Tầng Client: trình duyệt web (SPA) hoặc ứng dụng di động, giao tiếp với hệ thống qua HTTPS và WebSocket.")
B(doc, "Tầng API Gateway: điểm vào duy nhất, đảm nhiệm xác thực JWT, rate limit (sliding window), CORS, reverse proxy tới các vi dịch vụ backend, và duy trì WebSocket Hub cho kênh realtime.")
B(doc, "Tầng Vi dịch vụ: 6 service chính — Auth, Wallet, Market, Trading, Futures, Notification — mỗi service chạy trên cổng riêng, giao tiếp nội bộ qua gRPC.")
B(doc, "Tầng Hạ tầng: Redis (cache + event bus + WS fan-out), Kafka (durable event backbone), Elasticsearch (search + audit index).")
B(doc, "Tầng Lưu trữ: 6 instance PostgreSQL 16 độc lập, mỗi instance thuộc sở hữu của một vi dịch vụ duy nhất.")
IMG(doc, os.path.join(ASSETS, "fig1-architecture.png"), 16.0)
CAP(doc, "Hình 1. Kiến trúc tổng thể nền tảng Micro-Exchange")

H(doc, "4.2. Danh mục vi dịch vụ", 2)
TBL(doc,
    ["Vi dịch vụ", "Cổng", "Trách nhiệm chính", "CSDL riêng", "Giao tiếp"],
    [
        ["auth-service", "8081/9081", "Đăng ký, đăng nhập, JWT, 2FA TOTP, KYC, fraud, bonus, settings", "pg-auth :5551", "REST + gRPC"],
        ["wallet-service", "8082/9082", "Quản lý ví đa tiền tệ, nạp/rút, SePay webhook, khoá/mở khoá số dư", "pg-wallet :5552", "REST + gRPC"],
        ["market-service", "8083/9083", "Giá realtime Bybit WS, tổng hợp nến (1m–1w), tỷ giá VND/USDT", "pg-market :5553", "REST + gRPC"],
        ["trading-service", "8084", "Matching Engine (Spot), quản lý Order Book, đặt/huỷ lệnh", "pg-trading :5554", "REST"],
        ["futures-service", "8085", "Mở/đóng vị thế Futures, tính PnL, Liquidation Engine (tick 1s)", "pg-futures :5555", "REST"],
        ["notification-service", "8086", "Nhận domain event, tạo thông báo, đánh dấu đã đọc", "pg-notification :5556", "REST"],
        ["gateway", "8080", "Reverse proxy, JWT verify qua gRPC, rate limit, CORS, WS Hub", "—", "HTTP"],
        ["es-indexer", "—", "Consumer Kafka, index trades/orders vào Elasticsearch", "—", "Kafka consumer"],
    ],
    col_widths=[2.5, 2.0, 5.0, 3.0, 2.5],
)
CAP(doc, "Bảng 4.1. Danh mục vi dịch vụ và trách nhiệm")

H(doc, "4.3. Thiết kế cơ sở dữ liệu", 2)
P(doc, "Hình 8 trình bày sơ đồ quan hệ thực thể (ERD) tổng hợp. Do áp dụng nguyên tắc database-per-service, các quan hệ liên-CSDL (ví dụ users ↔ wallets) là quan hệ logic — được duy trì qua sự kiện user.registered chứ không có khoá ngoại vật lý.")
IMG(doc, os.path.join(ASSETS, "fig8-erd.png"), 16.0)
CAP(doc, "Hình 8. Sơ đồ quan hệ thực thể (ERD) tổng hợp")

TBL(doc,
    ["CSDL", "Cổng", "Vi dịch vụ", "Các bảng chính"],
    [
        ["pg-auth",         ":5551", "auth-service",         "users, kyc_profiles, kyc_documents, bonus_promotions, user_bonuses, fraud_logs, platform_settings"],
        ["pg-wallet",       ":5552", "wallet-service",       "wallets, deposits, withdrawals"],
        ["pg-market",       ":5553", "market-service",       "candles, user_trade_pairs"],
        ["pg-trading",      ":5554", "trading-service",      "orders, trades"],
        ["pg-futures",      ":5555", "futures-service",      "futures_positions"],
        ["pg-notification", ":5556", "notification-service", "notifications"],
    ],
    col_widths=[3.0, 1.5, 3.5, 8.0],
)
CAP(doc, "Bảng 4.2. Phân bổ CSDL PostgreSQL theo vi dịch vụ")

P(doc, "Một số bảng quan trọng và cấu trúc cột (trích từ GORM model):")
P(doc, "Bảng users: id (PK), email (UNIQUE), password_hash, full_name, phone, kyc_status (NONE/PENDING/VERIFIED/REJECTED), is_2fa, two_fa_secret (ẩn khỏi JSON), role (USER/ADMIN), email_verified, kyc_step (0–4), is_locked, lock_reason, last_login_ip, register_ip, created_at, updated_at.", bold=False, size=12)
P(doc, "Bảng wallets: id (PK), user_id + currency (UNIQUE INDEX), balance decimal(30,10), locked_balance decimal(30,10), updated_at.", size=12)
P(doc, "Bảng orders: id (PK), user_id, pair, side (BUY/SELL), type (LIMIT/MARKET/STOP_LIMIT), price decimal(30,10), stop_price, amount, filled_amount, status (OPEN/PARTIAL/FILLED/CANCELLED), created_at, updated_at.", size=12)
P(doc, "Bảng trades: id (PK), pair, buy_order_id, sell_order_id, buyer_id, seller_id, price, amount, total, buyer_fee, seller_fee, created_at.", size=12)
P(doc, "Bảng futures_positions: id (PK), user_id, pair, side (LONG/SHORT), leverage (1–125), entry_price, mark_price, size, margin, unrealized_pnl, liquidation_price, take_profit, stop_loss, status (OPEN/CLOSED/LIQUIDATED), created_at, closed_at.", size=12)

H(doc, "4.4. Thiết kế giao diện gRPC", 2)
P(doc, "Hệ thống định nghĩa 3 file .proto cho giao tiếp nội bộ:")
TBL(doc,
    ["Service gRPC", "File", "Phương thức", "Mô tả"],
    [
        ["AuthService",   "auth.proto",   "ValidateToken(token) → {valid, userId, email, role}", "Gateway gọi để xác thực JWT"],
        ["AuthService",   "auth.proto",   "GetUser(userId) → {id, email, name, role, kyc, 2fa}", "Lấy thông tin user"],
        ["WalletService", "wallet.proto", "CheckBalance(userId, currency, needed) → {sufficient, available}", "Kiểm tra số dư"],
        ["WalletService", "wallet.proto", "GetBalance(userId, currency) → {balance, locked}", "Đọc số dư và khoá"],
        ["WalletService", "wallet.proto", "Deduct(userId, currency, amount) → {newBalance}", "Trừ số dư"],
        ["WalletService", "wallet.proto", "Credit(userId, currency, amount) → {newBalance}", "Cộng số dư"],
        ["WalletService", "wallet.proto", "Lock(userId, currency, amount) → {success}", "Khoá số dư"],
        ["WalletService", "wallet.proto", "Unlock(userId, currency, amount) → {success}", "Mở khoá số dư"],
        ["MarketService", "market.proto", "GetPrice(pair) → {price}", "Lấy giá hiện tại"],
        ["MarketService", "market.proto", "GetAllTickers() → {tickers[]}", "Lấy toàn bộ bảng giá"],
    ],
    col_widths=[2.5, 2.5, 6.0, 5.0],
)
CAP(doc, "Bảng 4.3. Danh sách phương thức gRPC")

H(doc, "4.5. Thiết kế API REST", 2)
P(doc, "Bảng 4.4 liệt kê toàn bộ endpoint REST của hệ thống, tổ chức theo vi dịch vụ.")
TBL(doc,
    ["Phương thức", "Endpoint", "Auth?", "Mô tả"],
    [
        ["POST", "/api/auth/register", "Không", "Đăng ký tài khoản"],
        ["POST", "/api/auth/login", "Không", "Đăng nhập"],
        ["POST", "/api/auth/2fa/login", "Không", "Đăng nhập với mã TOTP"],
        ["POST", "/api/auth/refresh", "Không", "Refresh access token"],
        ["GET",  "/api/auth/profile", "Có", "Xem thông tin cá nhân"],
        ["PUT",  "/api/auth/profile", "Có", "Cập nhật thông tin"],
        ["POST", "/api/auth/2fa/enable", "Có", "Bật 2FA"],
        ["POST", "/api/auth/2fa/verify", "Có", "Xác nhận mã TOTP khi bật 2FA"],
        ["POST", "/api/kyc/email/send", "Có", "Gửi email xác thực"],
        ["POST", "/api/kyc/email/verify", "Có", "Xác thực mã email"],
        ["POST", "/api/kyc/profile", "Có", "Nộp hồ sơ KYC"],
        ["POST", "/api/kyc/document", "Có", "Tải lên giấy tờ KYC"],
        ["GET",  "/api/wallet/balances", "Có", "Xem tất cả số dư ví"],
        ["POST", "/api/wallet/deposit", "Có", "Tạo yêu cầu nạp tiền"],
        ["POST", "/api/wallet/withdraw", "Có", "Tạo yêu cầu rút tiền"],
        ["GET",  "/api/market/tickers", "Không", "Lấy bảng giá"],
        ["GET",  "/api/market/depth/:pair", "Không", "Lấy sổ lệnh"],
        ["GET",  "/api/market/candles/:pair", "Không", "Lấy biểu đồ nến"],
        ["POST", "/api/trading/orders", "Có", "Đặt lệnh Spot"],
        ["DELETE", "/api/trading/orders/:id", "Có", "Huỷ lệnh"],
        ["GET",  "/api/trading/orders", "Có", "Lịch sử lệnh"],
        ["POST", "/api/futures/order", "Có", "Mở vị thế Futures"],
        ["POST", "/api/futures/close/:id", "Có", "Đóng vị thế"],
        ["PUT",  "/api/futures/positions/:id/tpsl", "Có", "Cập nhật TP/SL"],
        ["GET",  "/api/notifications", "Có", "Danh sách thông báo"],
        ["POST", "/api/webhook/sepay", "Không", "Webhook SePay xác nhận nạp tiền"],
        ["GET",  "/ws", "Token", "WebSocket kết nối realtime"],
    ],
    col_widths=[2.0, 5.5, 1.5, 6.5],
)
CAP(doc, "Bảng 4.4. Danh sách API REST Endpoint")

H(doc, "4.6. Thiết kế Event Schema", 2)
TBL(doc,
    ["Topic", "Dữ liệu chính", "Producer", "Consumer tiêu biểu"],
    [
        ["trade.executed", "pair, buyerId, sellerId, price, amount, fee, side", "trading-service", "es-indexer, notification, wallet projector"],
        ["order.updated", "orderId, status, filledAmount, price", "trading-service", "DB projector, WebSocket"],
        ["balance.changed", "userId, currency, delta, reason, refId", "trading, wallet, futures", "wallet projector, audit"],
        ["position.changed", "positionId, userId, status, pnl, markPrice", "futures-service", "notification, futures projector"],
        ["user.registered", "userId, email, fullName", "auth-service", "wallet-service (tạo ví mặc định)"],
        ["notification.created", "userId, type, title, message", "notification-service", "WebSocket Hub"],
        ["price.updated", "pair, price, change24h, volume24h", "market-service", "trading, futures, WebSocket"],
    ],
    col_widths=[3.5, 5.0, 3.5, 4.0],
)
CAP(doc, "Bảng 4.5. Danh sách Domain Event")

H(doc, "4.7. Thiết kế kênh WebSocket", 2)
TBL(doc,
    ["Kênh", "Dữ liệu", "Tần suất", "Mô tả"],
    [
        ["trades@{pair}", "{price, amount, side}", "Mỗi giao dịch", "Giao dịch mới nhất của một cặp"],
        ["depth@{pair}", "{bids: [[price,qty]], asks: [[price,qty]]}", "Mỗi thay đổi sổ lệnh", "Snapshot 20 mức giá tốt nhất"],
        ["tickers@{pair}", "{price, change24h, volume24h}", "Mỗi cập nhật giá", "Bảng giá realtime"],
        ["candles@{pair}@{interval}", "{time, open, high, low, close, volume}", "Mỗi nến mới/update", "Dữ liệu cho biểu đồ TradingView"],
        ["notifications@{userId}", "{type, title, message}", "Khi có thông báo mới", "Thông báo cá nhân"],
        ["positions@{userId}", "{positionId, pnl, markPrice, status}", "Mỗi giây", "Cập nhật vị thế Futures"],
    ],
    col_widths=[3.5, 4.5, 3.0, 5.0],
)
CAP(doc, "Bảng 4.6. Danh sách kênh WebSocket")

H(doc, "4.8. Thuật toán khớp lệnh", 2)
P(doc, "Matching Engine là thành phần trung tâm của Trading Service, triển khai thuật toán khớp lệnh Price-Time Priority hoàn toàn trong bộ nhớ. Hình 10 minh hoạ cấu trúc dữ liệu Order Book và mã giả của thuật toán Match.")
IMG(doc, os.path.join(ASSETS, "fig10-orderbook.png"), 16.0)
CAP(doc, "Hình 10. Cấu trúc Order Book và thuật toán khớp lệnh Price-Time Priority")
P(doc, "Cấu trúc dữ liệu: Order Book cho mỗi cặp giao dịch bao gồm hai danh sách price level: Bids (sắp xếp giá giảm dần — giá tốt nhất ở đầu) và Asks (sắp xếp giá tăng dần). Mỗi price level chứa một danh sách lệnh theo thứ tự thời gian (FIFO). Khi lệnh mới đến:")
B(doc, "Nếu lệnh BUY: so sánh với Asks[0] (giá bán thấp nhất). MARKET luôn khớp; LIMIT chỉ khớp khi best ask ≤ giá lệnh mua.")
B(doc, "Nếu lệnh SELL: so sánh với Bids[0] (giá mua cao nhất). MARKET luôn khớp; LIMIT chỉ khớp khi best bid ≥ giá lệnh bán.")
B(doc, "Số lượng khớp = min(remaining_incoming, remaining_resting). Cập nhật filledAmount cả hai phía.")
B(doc, "Lặp cho đến khi incoming hết số lượng hoặc không còn mức giá thoả mãn.")
B(doc, "Phần chưa khớp của lệnh LIMIT được thêm vào sổ lệnh chờ.")

H(doc, "4.9. Thuật toán thanh lý vị thế Futures", 2)
P(doc, "Liquidation Engine chạy trong một goroutine riêng, thực hiện kiểm tra PnL mỗi 1 giây cho tất cả vị thế đang mở. Công thức tính PnL chưa thực hiện (Unrealized PnL):")
P(doc, "LONG:  PnL = Size × (MarkPrice − EntryPrice)", indent=False, bold=True, size=12)
P(doc, "SHORT: PnL = Size × (EntryPrice − MarkPrice)", indent=False, bold=True, size=12)
P(doc, "Công thức tính giá thanh lý (Liquidation Price):")
P(doc, "LONG:  LiqPrice = EntryPrice × (1 − 1/Leverage + 0.005)", indent=False, bold=True, size=12)
P(doc, "SHORT: LiqPrice = EntryPrice × (1 + 1/Leverage − 0.005)", indent=False, bold=True, size=12)
P(doc, "Trong đó, hằng số 0.005 (0,5%) đại diện cho tỷ lệ ký quỹ duy trì (maintenance margin ratio). Khi PnL ≤ −(Margin × 0.995), vị thế bị đánh dấu LIQUIDATED, hệ thống hoàn trả phần margin còn lại (nếu có) cho người dùng qua gRPC Credit.")

H(doc, "4.10. Thiết kế Redis Lua Script cho số dư", 2)
P(doc, "Để đảm bảo tính nguyên tử của các thao tác số dư mà không cần lock phân tán, hệ thống sử dụng Redis Lua script. Mỗi script chạy nguyên tử (atomic) trên Redis server — không có thao tác nào xen giữa. Các script chính:")
B(doc, "BalanceDeduct: Kiểm tra bal ≥ amount → trừ → trả về số dư mới. Nếu không đủ trả về -1, nếu key không tồn tại trả về -2.")
B(doc, "BalanceCredit: Cộng amount vào bal. Nếu key chưa tồn tại, khởi tạo bằng amount.")
B(doc, "BalanceLock: Kiểm tra (bal − locked) ≥ amount → tăng locked. Đảm bảo không khoá quá số dư khả dụng.")
B(doc, "BalanceUnlock: Giảm locked, clamp xuống 0 nếu âm.")
P(doc, "Key naming convention: bal:{userID}:{currency} cho số dư, locked:{userID}:{currency} cho số khoá. Ví dụ: bal:42:USDT, locked:42:BTC.")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 5. CONG NGHE
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 5. NGĂN XẾP CÔNG NGHỆ", 1)

H(doc, "5.1. Bảng công nghệ chi tiết", 2)
TBL(doc,
    ["Lớp", "Công nghệ / Thư viện", "Phiên bản", "Vai trò"],
    [
        ["Ngôn ngữ",      "Go (Golang)",            "1.22+",      "Ngôn ngữ chính cho toàn bộ backend"],
        ["Web framework",  "Gin",                    "v1",         "HTTP router, middleware cho từng service"],
        ["ORM",            "GORM",                   "v2",         "Mapping struct ↔ PostgreSQL, auto-migrate"],
        ["RPC",            "gRPC + Protocol Buffers", "v1",        "Giao tiếp đồng bộ giữa các service"],
        ["CSDL quan hệ",   "PostgreSQL",             "16-alpine", "Lưu trữ bền vững: user, wallet, order, trade, position"],
        ["In-memory cache", "Redis",                  "7-alpine",  "Cache số dư, Lua atomic, Pub/Sub WS, Streams event bus"],
        ["Event broker",   "Apache Kafka",            "3.9 (KRaft)","Backbone sự kiện bền vững, hỗ trợ replay"],
        ["Search engine",  "Elasticsearch",           "8.17",      "Index giao dịch, audit, truy vấn lịch sử"],
        ["Realtime",       "Gorilla WebSocket",       "v1",        "Kênh WS full-duplex tới client"],
        ["Xác thực",       "JWT (HS256)",             "—",         "Token xác thực do Auth Service phát hành"],
        ["2FA",            "TOTP (RFC 6238)",         "—",         "Mã OTP 6 chữ số, tương thích Google Authenticator"],
        ["Mật khẩu",       "bcrypt",                  "cost=10",   "Hàm băm chậm chống brute-force"],
        ["Giá real-time",  "Bybit WebSocket",         "—",         "Nguồn feed giá crypto real-time"],
        ["Tỷ giá forex",   "CoinGecko API",           "v3",        "Fallback giá, tỷ giá forex"],
        ["Thanh toán",     "SePay",                   "—",         "QR chuyển khoản ngân hàng VND"],
        ["Container",      "Docker Compose",           "—",         "Điều phối hạ tầng phát triển và kiểm thử"],
    ],
    col_widths=[2.5, 4.0, 2.0, 7.5],
)
CAP(doc, "Bảng 5.1. Ngăn xếp công nghệ chi tiết")

H(doc, "5.2. Lý do lựa chọn", 2)
B(doc, "Go: được chọn làm ngôn ngữ chính nhờ tính năng concurrency sẵn có (goroutine, channel), garbage collector độ trễ thấp (P99 < 1ms), biên dịch thành binary tĩnh duy nhất, và hiệu năng ngang C++ cho các bài toán I/O-bound. Đây cũng là ngôn ngữ phổ biến nhất trong ngành exchange (Binance, Bybit đều sử dụng Go).")
B(doc, "Redis Lua: thay vì dùng lock phân tán (Redlock) hay transaction DB, Lua script cho phép bundle nhiều thao tác thành một đơn vị nguyên tử trên Redis — đơn giản hơn, nhanh hơn, và đủ an toàn cho use case số dư.")
B(doc, "PostgreSQL 16: hỗ trợ kiểu decimal(30,10) và ràng buộc mạnh, phù hợp cho tính toán tiền tệ. Phiên bản 16 cải thiện đáng kể hiệu năng parallel query.")
B(doc, "Kafka KRaft: loại bỏ dependency ZooKeeper, đơn giản hoá triển khai. Kafka đảm bảo ordering per-partition và cho phép consumer replay sự kiện.")
B(doc, "Elasticsearch 8.17: hỗ trợ full-text search, aggregation và time-series query, phù hợp cho lịch sử giao dịch và báo cáo.")

H(doc, "5.3. Cấu trúc thư mục dự án", 2)
P(doc, "Dự án sử dụng Go Workspace (go.work) để quản lý nhiều module cùng lúc. Cấu trúc thư mục tuân thủ Standard Go Project Layout:", indent=False)
P(doc,
    "micro-exchange/\n"
    "├── go.work                  # Go workspace\n"
    "├── docker-compose.yml       # Hạ tầng (6 PG + Redis + Kafka + ES)\n"
    "├── shared/                  # Thư viện dùng chung\n"
    "│   ├── model/               # GORM models (user, wallet, order, ...)\n"
    "│   ├── proto/               # gRPC proto files + generated code\n"
    "│   ├── eventbus/            # Redis Streams event bus\n"
    "│   ├── redisutil/           # Balance cache + Lua scripts + Rate limiter\n"
    "│   ├── middleware/          # JWT middleware, CORS\n"
    "│   ├── ws/                  # WebSocket Hub\n"
    "│   └── response/            # Chuẩn hoá HTTP response\n"
    "├── auth-service/            # Cmd + internal (handler, service, repo)\n"
    "├── wallet-service/\n"
    "├── market-service/\n"
    "├── trading-service/         # + internal/engine (OrderBook, Match)\n"
    "├── futures-service/\n"
    "├── notification-service/\n"
    "├── gateway/\n"
    "└── es-indexer/",
    indent=False, size=11)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 6. HIEN THUC
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 6. HIỆN THỰC HỆ THỐNG", 1)

H(doc, "6.1. Auth Service", 2)
P(doc, "Auth Service là dịch vụ chịu trách nhiệm toàn bộ vòng đời xác thực và quản lý người dùng. Các thành phần chính:")
B(doc, "AuthService: xử lý đăng ký (bcrypt hash mật khẩu), đăng nhập (phát JWT HS256), refresh token, đổi mật khẩu. Khi đăng ký thành công, phát sự kiện user.registered để Wallet Service tạo ví mặc định.")
B(doc, "TOTPService: sinh secret TOTP tương thích RFC 6238, tạo QR code URI cho Google Authenticator, xác thực mã 6 chữ số.")
B(doc, "KYC Handler: quản lý quy trình 4 bước — (1) xác thực email, (2) điền hồ sơ cá nhân (KYCProfile), (3) tải lên giấy tờ (CCCD mặt trước, mặt sau, selfie), (4) Admin duyệt. Mỗi bước cập nhật kyc_step trên bảng users.")
B(doc, "Admin Handler: CRUD người dùng, duyệt/từ chối KYC, thống kê (tổng user, tổng deposit, tổng trade), biểu đồ tăng trưởng.")
B(doc, "Fraud Handler: ghi và quản lý nhật ký gian lận — FraudLog với các loại: BONUS_FARMING, WASH_TRADING, MULTI_ACCOUNT.")
B(doc, "Bonus Handler: quản lý chương trình khuyến mãi (BonusPromotion), tự động cộng bonus khi người dùng nạp tiền (trigger ON_DEPOSIT).")
B(doc, "Settings Handler: cấu hình platform_settings (phí giao dịch, cặp coin hoạt động, giới hạn rút tiền).")

H(doc, "6.2. Wallet Service", 2)
P(doc, "Wallet Service quản lý ví đa tiền tệ cho mỗi người dùng. Khi nhận sự kiện user.registered, service tự động tạo ví VND, USDT và tất cả coin đang hoạt động.")
B(doc, "Nạp tiền VND: tạo mã đơn (orderCode) độc nhất, sinh QR SePay chứa nội dung chuyển khoản. Khi nhận webhook POST /api/webhook/sepay, xác thực chữ ký, cộng số dư VND, quy đổi sang USDT theo tỷ giá hiện hành, phát sự kiện balance.changed.")
B(doc, "Rút tiền: người dùng gửi yêu cầu với ngân hàng, số tài khoản, số tiền. Hệ thống khoá số dư tương ứng, tạo Withdrawal status PENDING. Admin duyệt → APPROVED hoặc từ chối → REJECTED (mở khoá số dư).")
B(doc, "gRPC Server: cung cấp 7 phương thức cho các service khác gọi: CheckBalance, GetBalance, Deduct, Credit, Lock, Unlock, UpdateBalance.")
B(doc, "Balance Cache: mọi thao tác số dư đều đi qua BalanceCache (Redis Lua), sau đó phát sự kiện để projector ghi vào PostgreSQL.")

H(doc, "6.3. Market Service", 2)
P(doc, "Market Service chịu trách nhiệm thu thập và phân phối dữ liệu giá thị trường.")
B(doc, "PriceFeed: kết nối Bybit WebSocket v5, subscribe ticker của 55 cặp giao dịch. Khi nhận giá mới, cập nhật Redis key price:{pair} và broadcast qua WebSocket tickers@{pair}.")
B(doc, "CandleAggregator: tổng hợp nến từ tick giá realtime cho 7 khung thời gian (1m, 5m, 15m, 1h, 4h, 1d, 1w). Mỗi nến được lưu vào PostgreSQL và broadcast qua WS candles@{pair}@{interval}.")
B(doc, "CandleBackfill: khi service khởi động, lấy lịch sử nến từ Bybit REST API để điền dữ liệu quá khứ.")
B(doc, "VND Rate: lấy tỷ giá VND/USDT từ nguồn bên ngoài, cache trong Redis, refresh mỗi 5 phút.")

H(doc, "6.4. Trading Service và Matching Engine", 2)
P(doc, "Trading Service là trung tâm giao dịch Spot, bao gồm hai thành phần chính:")
B(doc, "MatchingEngine: quản lý map[pair]*OrderBook + map[pair]*Mutex. Mỗi cặp giao dịch có riêng một order book trong bộ nhớ và một mutex để serialize việc khớp lệnh. Khi ProcessOrder được gọi, engine thực hiện: (1) Match trong bộ nhớ, (2) cập nhật số dư qua Redis Lua cho từng trade, (3) unlock số dư đã khoá, (4) phát sự kiện, (5) broadcast WS.")
B(doc, "OrderBook: triển khai bằng hai Go slice — Bids (giảm dần) và Asks (tăng dần). Mỗi PriceLevel chứa danh sách *Order theo FIFO. Phương thức Match() trả về []TradeResult. Cold-start recovery: khi khởi động, LoadOpenOrders() tải tất cả lệnh OPEN/PARTIAL từ DB vào sổ lệnh.")
B(doc, "OrderService: ghi và đọc lệnh từ PostgreSQL, xử lý projector cập nhật trạng thái lệnh khi nhận sự kiện order.updated.")

H(doc, "6.5. Futures Service và Liquidation Engine", 2)
B(doc, "FuturesService: xử lý mở vị thế (OpenPosition), đóng thủ công (ClosePosition), cập nhật TP/SL. Khi mở vị thế, gọi gRPC CheckBalance + Deduct, tính toán margin, fee, liquidation price, rồi ghi vào DB.")
B(doc, "LiquidationEngine: goroutine chạy mỗi 1 giây — lấy giá mark từ Redis, load tất cả vị thế OPEN, tính PnL cho từng vị thế. Nếu PnL vượt ngưỡng thanh lý: (1) cập nhật status = LIQUIDATED, (2) gRPC Credit phần margin còn lại, (3) phát sự kiện position.changed + balance.changed, (4) gửi notification POSITION_LIQUIDATED.")

H(doc, "6.6. Notification Service", 2)
P(doc, "Notification Service subscribe các sự kiện: trade.executed (thông báo lệnh khớp), position.changed (thông báo mở/đóng/thanh lý vị thế), balance.changed khi lý do là deposit (thông báo nạp tiền thành công). Mỗi sự kiện sinh một bản ghi Notification trong pg-notification và broadcast qua WebSocket notifications@{userId}.")

H(doc, "6.7. API Gateway", 2)
P(doc, "Gateway được xây dựng trên Gin framework, hoạt động như reverse proxy thông minh:")
B(doc, "Route table: ánh xạ path prefix → target URL + boolean yêu cầu xác thực. Sử dụng longest-prefix matching.")
B(doc, "Xác thực: nếu route yêu cầu auth, đọc JWT từ cookie access_token hoặc header Authorization: Bearer. Gọi gRPC ValidateToken tới Auth Service. Nếu hợp lệ, inject header X-User-ID, X-User-Email, X-User-Role vào request trước khi proxy.")
B(doc, "Rate Limiter: sliding window trên Redis (Lua script), giới hạn 5 req/phút cho login, 1000 req/phút cho global.")
B(doc, "WebSocket Hub: gateway duy trì kết nối WS với client, subscribe kênh Redis Pub/Sub ws:broadcast để nhận data từ backend services và fan-out tới client.")

H(doc, "6.8. ES Indexer", 2)
P(doc, "ES Indexer là worker service không có API endpoint, chỉ consume sự kiện từ Kafka. Khi nhận trade.executed, tạo document TradeDoc và index vào Elasticsearch index \"trades\". Tương tự cho order.updated → index \"orders\". ESClient wrapper đảm bảo index được tạo với mapping phù hợp (keyword cho pair, date cho createdAt).")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 7. LUONG NGHIEP VU
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 7. LUỒNG NGHIỆP VỤ CHI TIẾT", 1)

H(doc, "7.1. Luồng đặt lệnh Spot", 2)
P(doc, "Hình 2 trình bày luồng dữ liệu (data flow) 15 bước của quá trình đặt và khớp lệnh Spot. Hình 6 bổ sung sơ đồ tuần tự (sequence diagram) chi tiết, thể hiện thứ tự tương tác giữa các đối tượng.")
IMG(doc, os.path.join(ASSETS, "fig2-spot-flow.png"), 16.0)
CAP(doc, "Hình 2. Luồng xử lý lệnh Spot theo mô hình CQRS")
IMG(doc, os.path.join(ASSETS, "fig6-seq-spot.png"), 16.0)
CAP(doc, "Hình 6. Sơ đồ tuần tự — Đặt lệnh giao dịch Spot")
P(doc, "Điểm nổi bật: toàn bộ quá trình khớp lệnh và cập nhật số dư (bước 6–8) chỉ thao tác trên bộ nhớ và Redis — không có disk I/O trên đường critical path. CSDL PostgreSQL chỉ được ghi bất đồng bộ bởi projector (bước 11) sau khi nhận sự kiện từ Redis Streams/Kafka. Điều này đảm bảo hot path có độ trễ cực thấp.")

H(doc, "7.2. Luồng thanh lý vị thế Futures", 2)
P(doc, "Hình 3 trình bày luồng dữ liệu mở và thanh lý vị thế Futures. Hình 7 bổ sung sơ đồ tuần tự của Liquidation Engine.")
IMG(doc, os.path.join(ASSETS, "fig3-futures-flow.png"), 16.0)
CAP(doc, "Hình 3. Luồng mở và thanh lý vị thế Futures")
IMG(doc, os.path.join(ASSETS, "fig7-seq-liquidation.png"), 16.0)
CAP(doc, "Hình 7. Sơ đồ tuần tự — Thanh lý vị thế Futures")
P(doc, "Liquidation Engine là goroutine chạy liên tục, kiểm tra tất cả vị thế OPEN mỗi giây. Khi phát hiện PnL vượt ngưỡng, engine đóng vị thế cưỡng chế, hoàn trả margin còn lại và gửi thông báo. Cơ chế này tương tự cách Binance và Bybit thực hiện auto-deleverage.")

H(doc, "7.3. Luồng nạp tiền VND qua SePay", 2)
B(doc, "Bước 1: Người dùng gọi POST /api/wallet/deposit {amount, currency: VND}.")
B(doc, "Bước 2: Wallet Service tạo Deposit record (status PENDING), sinh orderCode duy nhất.")
B(doc, "Bước 3: Trả về QR SePay chứa nội dung chuyển khoản \"ME-{orderCode}\".")
B(doc, "Bước 4: Người dùng mở app ngân hàng, quét QR, chuyển khoản.")
B(doc, "Bước 5: SePay gọi webhook POST /api/webhook/sepay — nội dung giao dịch chứa mã ME-{orderCode}.")
B(doc, "Bước 6: Wallet Service xác thực webhook, tìm Deposit theo orderCode, cập nhật status CONFIRMED.")
B(doc, "Bước 7: Cộng số dư VND và quy đổi sang USDT (theo tỷ giá VND/USDT hiện tại), phát sự kiện balance.changed.")
B(doc, "Bước 8: Notification Service gửi thông báo DEPOSIT_CONFIRMED tới client.")

H(doc, "7.4. Luồng rút tiền", 2)
B(doc, "Bước 1: Người dùng gọi POST /api/wallet/withdraw {amount, bankCode, bankAccount, accountName}.")
B(doc, "Bước 2: Wallet Service kiểm tra số dư khả dụng, khoá số dư tương ứng.")
B(doc, "Bước 3: Tạo Withdrawal record status PENDING.")
B(doc, "Bước 4: Admin xem danh sách yêu cầu rút tiền, duyệt (APPROVED) hoặc từ chối (REJECTED).")
B(doc, "Bước 5: Nếu APPROVED — trừ số dư vĩnh viễn, phát balance.changed. Nếu REJECTED — mở khoá số dư.")

H(doc, "7.5. Luồng xác thực và KYC", 2)
P(doc, "Quy trình đăng ký và KYC gồm các bước tuần tự:")
B(doc, "Bước 1: Đăng ký — POST /api/auth/register {email, password, fullName}. Hệ thống hash mật khẩu bcrypt, tạo user, phát user.registered.")
B(doc, "Bước 2: Đăng nhập — POST /api/auth/login. Nếu 2FA bật → trả yêu cầu mã TOTP → POST /api/auth/2fa/login.")
B(doc, "Bước 3: KYC Bước 1 — Gửi mã xác thực email (POST /api/kyc/email/send), xác nhận mã (POST /api/kyc/email/verify). kycStep = 1.")
B(doc, "Bước 4: KYC Bước 2 — Nộp hồ sơ cá nhân (POST /api/kyc/profile): họ tên, ngày sinh, địa chỉ, nghề nghiệp, thu nhập, kinh nghiệm giao dịch. kycStep = 2.")
B(doc, "Bước 5: KYC Bước 3 — Tải lên giấy tờ (POST /api/kyc/document): CCCD mặt trước, mặt sau, ảnh selfie. kycStep = 3.")
B(doc, "Bước 6: KYC Bước 4 — Admin duyệt (POST /api/admin/kyc/:userId/approve). kycStatus = VERIFIED, kycStep = 4.")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 8. EVENT-DRIVEN + CQRS
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 8. CƠ CHẾ EVENT-DRIVEN VÀ CQRS", 1)

H(doc, "8.1. Kiến trúc Event Bus", 2)
P(doc, "Hình 4 minh hoạ kiến trúc fan-out của event bus: bên trái là 5 producer service, giữa là Event Backbone (Redis Streams + Kafka), bên phải là 5 consumer. Mỗi consumer hoạt động trong một consumer group độc lập, đảm bảo mỗi sự kiện được xử lý đúng một lần trong nhóm nhưng tất cả nhóm đều nhận bản sao.")
IMG(doc, os.path.join(ASSETS, "fig4-event-bus.png"), 16.0)
CAP(doc, "Hình 4. Sơ đồ Event Bus — Producer · Broker · Consumer")
P(doc, "Redis Streams được chọn làm hot-path bus vì độ trễ sub-millisecond, phù hợp cho việc đồng bộ số dư và WebSocket broadcast trong real-time. Kafka được dùng song song cho các consumer cần tính bền vững (durable): es-indexer cần replay sự kiện khi rebuild index, analytics cần lưu trữ lâu dài.")
P(doc, "Cách thức hoạt động: producer gọi Bus.Publish() → ghi vào Redis Stream (XADD, maxlen ~50.000 entry). Consumer đăng ký consumer group (XREADGROUP), xử lý và ACK. Nếu handler lỗi → không ACK → sự kiện được trao lại (pending entries list). Kafka adapter làm cầu nối: consume từ Redis Streams và produce sang Kafka topic tương ứng.")

H(doc, "8.2. Mô hình CQRS trên hot path", 2)
P(doc, "Trong Micro-Exchange, CQRS được hiện thực hoá rõ nhất tại Trading Service:")
B(doc, "Command side (Write): Client → Gateway → Trading Handler → MatchingEngine.ProcessOrder() → OrderBook.Match() (in-memory) → Redis Lua updateBalance() → Bus.Publish().")
B(doc, "Query side (Read): Client → Gateway → Trading Handler → OrderService.GetHistory() → PostgreSQL pg-trading → trả JSON. Dữ liệu trong PostgreSQL được cập nhật bởi projector consumer tiêu thụ order.updated, trade.executed.")
P(doc, "Lợi ích: command path không chờ disk I/O, đạt độ trễ < 5ms. Query path có thể tối ưu index riêng (composite index trên pair + status, index trên created_at) mà không ảnh hưởng đến hiệu năng ghi.")

H(doc, "8.3. Danh sách Domain Event", 2)
P(doc, "Xem Bảng 4.5 (Chương 4) — liệt kê 7 loại domain event chính. Mỗi event là một thông điệp bất biến (immutable fact) mô tả sự thật đã xảy ra, được serialize dưới dạng JSON và lưu vào Redis Stream với metadata: topic, data, timestamp (unix ms).")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 9. BAO MAT
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 9. BẢO MẬT", 1)

H(doc, "9.1. Xác thực và phân quyền", 2)
B(doc, "Mật khẩu: băm bằng bcrypt với cost = 10 trước khi lưu DB. Không bao giờ lưu hoặc log cleartext.")
B(doc, "JWT: ký bằng HS256, thời gian sống ngắn (15 phút), refresh token 7 ngày. Gateway xác thực qua gRPC call đồng bộ tới Auth Service.")
B(doc, "2FA TOTP: secret lưu dạng base32 trong cột two_fa_secret (json:\"-\"), không lộ qua API. Khi bật 2FA, đăng nhập yêu cầu mã 6 chữ số từ Google Authenticator.")
B(doc, "Phân quyền: middleware kiểm tra header X-User-Role. Route /api/admin/* yêu cầu role = ADMIN. Route giao dịch yêu cầu KYC đạt bước tối thiểu.")

H(doc, "9.2. Bảo vệ API", 2)
B(doc, "Rate Limit: sliding window trên Redis (Lua script). Login: 10 req/phút. Register: 5 req/phút. Global: 1000 req/phút. Trả HTTP 429 khi vượt ngưỡng.")
B(doc, "CORS: chỉ cho phép origin cấu hình, đặt tại middleware.CORS().")
B(doc, "Input validation: Gin binding tag (required, min, max, oneof) kiểm tra đầu vào tại handler, trước khi dữ liệu đi vào service layer.")
B(doc, "CSRF: sử dụng Bearer token trong header (không dùng cookie cho API), giảm thiểu rủi ro CSRF.")

H(doc, "9.3. Chống gian lận", 2)
P(doc, "Auth Service tích hợp module FraudService ghi nhận hành vi nghi ngờ:")
B(doc, "BONUS_FARMING: nhiều tài khoản từ cùng IP đăng ký liên tục để nhận bonus.")
B(doc, "WASH_TRADING: hai tài khoản tự mua bán qua lại để tạo khối lượng giả.")
B(doc, "MULTI_ACCOUNT: phát hiện nhiều tài khoản chia sẻ IP/fingerprint.")
P(doc, "Mỗi FraudLog lưu: userIDs liên quan, loại gian lận, bằng chứng (JSON), hành động đã thực hiện (FLAGGED, ACCOUNTS_LOCKED, BONUS_REVOKED, DISMISSED). Admin có thể review và lock tài khoản nghi vấn.")

H(doc, "9.4. Ánh xạ OWASP Top 10", 2)
TBL(doc,
    ["OWASP", "Rủi ro", "Biện pháp trong Micro-Exchange"],
    [
        ["A01", "Broken Access Control", "JWT + role check middleware; route-level auth; admin gate"],
        ["A02", "Cryptographic Failures", "bcrypt hash; HS256 JWT; TLS (khi deploy production)"],
        ["A03", "Injection", "GORM parameterized query; Gin binding validation; Lua script"],
        ["A04", "Insecure Design", "CQRS tách hot path; per-pair mutex; database-per-service"],
        ["A05", "Security Misconfiguration", "CORS whitelist; healthcheck endpoints; Docker non-root"],
        ["A06", "Vulnerable Components", "Go module + sum verification; Docker alpine images"],
        ["A07", "Auth Failures", "Rate limit login; 2FA TOTP; account lock after fraud flag"],
        ["A08", "Data Integrity Failures", "Lua atomic scripts; event sourcing; KYC step validation"],
        ["A09", "Logging & Monitoring", "Structured log; fraud_logs table; consumer lag monitoring"],
        ["A10", "Server-Side Request Forgery", "Internal gRPC only; no user-controlled URL fetch"],
    ],
    col_widths=[1.5, 4.0, 10.5],
)
CAP(doc, "Bảng 9.1. Ánh xạ OWASP Top 10 với biện pháp bảo mật")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 10. TRIEN KHAI
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 10. TRIỂN KHAI VÀ KHẢ NĂNG MỞ RỘNG", 1)

H(doc, "10.1. Sơ đồ triển khai Docker Compose", 2)
P(doc, "Hình 9 trình bày sơ đồ triển khai. Toàn bộ hạ tầng (6 PostgreSQL, Redis, Kafka, Elasticsearch) chạy trên Docker Compose. Các Go service chạy trực tiếp trên host (hoặc đóng gói Dockerfile riêng). Lệnh khởi động:")
P(doc, "docker compose up -d   # Hạ tầng\n./start-all.sh           # 8 Go service song song", indent=False, bold=True, size=12)
IMG(doc, os.path.join(ASSETS, "fig9-deployment.png"), 16.0)
CAP(doc, "Hình 9. Sơ đồ triển khai Docker Compose")

H(doc, "10.2. Biến môi trường", 2)
TBL(doc,
    ["Biến", "Giá trị mặc định", "Mô tả"],
    [
        ["PORT", "8080 (gateway)", "Cổng HTTP của từng service"],
        ["JWT_SECRET", "(bắt buộc)", "Khoá bí mật ký JWT HS256"],
        ["REDIS_URL", "redis://localhost:6379", "Kết nối Redis"],
        ["KAFKA_BROKERS", "localhost:9192", "Địa chỉ Kafka broker"],
        ["ES_URL", "http://localhost:9201", "Địa chỉ Elasticsearch"],
        ["DB_DSN", "postgres://postgres:postgres@localhost:555X/...", "DSN PostgreSQL cho mỗi service"],
        ["AUTH_URL", "http://localhost:8081", "URL Auth Service (cho gateway proxy)"],
        ["WALLET_URL", "http://localhost:8082", "URL Wallet Service"],
        ["MARKET_URL", "http://localhost:8083", "URL Market Service"],
        ["SEPAY_BANK_CODE", "(cấu hình)", "Mã ngân hàng tích hợp SePay"],
    ],
    col_widths=[3.5, 5.0, 7.5],
)
CAP(doc, "Bảng 10.1. Biến môi trường chính")

H(doc, "10.3. Khả năng mở rộng", 2)
B(doc, "Horizontal scale per-service: mỗi service có thể chạy nhiều instance, Gateway load balance qua reverse proxy.")
B(doc, "Per-pair sharding cho Trading Service: mỗi pair có mutex riêng, có thể phân tán khác pair sang instance khác.")
B(doc, "Per-user sharding cho Futures Service: Liquidation Engine có thể chia theo userId hoặc pair.")
B(doc, "Read replica PostgreSQL: CSDL có thể cấu hình replica read-only phục vụ báo cáo, giảm tải cho primary.")
B(doc, "Kafka partition: tăng số partition cho topic có throughput cao, mỗi partition xử lý bởi một consumer trong group.")
B(doc, "Elasticsearch cluster: thêm node cho tìm kiếm phân tán khi dữ liệu lớn.")
B(doc, "Redis Cluster: chuyển từ single Redis sang cluster mode khi cần > 100GB RAM hoặc > 1M ops/s.")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# CHUONG 11. KET LUAN
# ═══════════════════════════════════════════════════════════════════
H(doc, "CHƯƠNG 11. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)

H(doc, "11.1. Kết quả đạt được", 2)
P(doc, "Đồ án đã hiện thực hoá thành công một nền tảng giao dịch tiền mã hoá đầy đủ chức năng theo kiến trúc vi dịch vụ, bao gồm:")
B(doc, "8 vi dịch vụ hoạt động độc lập: Auth, Wallet, Market, Trading, Futures, Notification, Gateway, ES Indexer.")
B(doc, "6 CSDL PostgreSQL riêng biệt, tuân thủ nguyên tắc database-per-service.")
B(doc, "Matching Engine khớp lệnh Spot hoàn toàn trong bộ nhớ, hot path không chạm DB.")
B(doc, "Liquidation Engine kiểm tra PnL mỗi giây, thanh lý tự động với đòn bẩy lên đến 125x.")
B(doc, "Event bus kết hợp Redis Streams + Kafka, hỗ trợ fan-out tới 5 loại consumer.")
B(doc, "55 cặp giao dịch (50 crypto + 3 forex + 2 commodity) với giá realtime từ Bybit.")
B(doc, "Hệ thống bảo mật đa lớp: bcrypt, JWT, 2FA TOTP, rate limit, KYC 4 bước, fraud detection.")
B(doc, "Tích hợp thanh toán SePay cho nạp VND qua QR ngân hàng.")

H(doc, "11.2. Hạn chế", 2)
B(doc, "Sử dụng float64 cho kiểu tiền tệ — có nguy cơ sai số lũy tích khi khối lượng giao dịch lớn. Cần chuyển sang decimal cố định (fixed-point) hoặc thư viện shopspring/decimal.")
B(doc, "Rollback của Matching Engine khi Redis gặp sự cố giữa chừng chưa hoàn thiện — có thể dẫn đến duplicate entry trong order book.")
B(doc, "Chưa có outbox pattern — nếu process crash giữa Redis commit và event publish, dữ liệu sẽ lệch giữa Redis và DB.")
B(doc, "Chưa có bộ kiểm thử tự động (unit test, integration test) cho Matching Engine — đây là thành phần quan trọng nhất cần coverage cao.")
B(doc, "Chưa có giao diện frontend — toàn bộ tương tác qua API (Postman/curl/WS client).")

H(doc, "11.3. Hướng phát triển", 2)
B(doc, "Chuyển đổi kiểu tiền tệ sang decimal cố định (ưu tiên cao nhất cho production).")
B(doc, "Áp dụng outbox pattern: ghi sự kiện vào bảng outbox cùng transaction với thao tác balance, sau đó một worker poll outbox và publish lên Kafka.")
B(doc, "Bổ sung sản phẩm Options (quyền chọn), chỉ báo kỹ thuật (MA, RSI, MACD) trên biểu đồ.")
B(doc, "Tích hợp risk engine với mô hình VaR (Value at Risk) để đánh giá rủi ro danh mục.")
B(doc, "Xây dựng frontend SPA (React/Next.js) với TradingView charting library.")
B(doc, "Tích hợp thêm cổng thanh toán quốc tế (Stripe, MoonPay) cho nạp tiền bằng thẻ quốc tế.")
B(doc, "Triển khai trên Kubernetes với Helm chart, auto-scaling và monitoring (Prometheus + Grafana).")
B(doc, "Viết bộ test toàn diện: unit test cho matching engine, integration test cho luồng nạp/rút, load test cho hot path.")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# PHU LUC A. DANH MUC TU VIET TAT VA THUAT NGU
# ═══════════════════════════════════════════════════════════════════
H(doc, "PHỤ LỤC A. DANH MỤC TỪ VIẾT TẮT VÀ THUẬT NGỮ", 1)
TBL(doc,
    ["Viết tắt / Thuật ngữ", "Đầy đủ", "Giải thích"],
    [
        ["API",       "Application Programming Interface",    "Giao diện lập trình ứng dụng giữa các thành phần phần mềm."],
        ["REST",      "Representational State Transfer",      "Kiểu kiến trúc giao tiếp qua HTTP theo mô hình tài nguyên."],
        ["JWT",       "JSON Web Token",                       "Token xác thực dạng JSON có chữ ký số, RFC 7519."],
        ["gRPC",      "gRPC Remote Procedure Call",           "Giao thức RPC hiệu năng cao, sử dụng HTTP/2 + Protobuf."],
        ["Protobuf",  "Protocol Buffers",                     "Định dạng mã hoá nhị phân cấu trúc dữ liệu, dùng cho gRPC."],
        ["ORM",       "Object-Relational Mapping",            "Kỹ thuật ánh xạ bảng CSDL sang đối tượng trong ngôn ngữ lập trình."],
        ["CQRS",      "Command Query Responsibility Segregation", "Tách đường ghi (command) và đọc (query) thành các mô hình riêng."],
        ["EDA",       "Event-Driven Architecture",            "Kiến trúc phần mềm dựa trên sự kiện làm trung tâm truyền thông."],
        ["KYC",       "Know Your Customer",                   "Quy trình xác minh danh tính khách hàng."],
        ["AML",       "Anti-Money Laundering",                "Các biện pháp phòng chống rửa tiền."],
        ["2FA",       "Two-Factor Authentication",            "Xác thực hai yếu tố."],
        ["TOTP",      "Time-based One-Time Password",         "Mật khẩu dùng một lần dựa trên thời gian, RFC 6238."],
        ["PnL",       "Profit and Loss",                      "Lãi/lỗ của một vị thế giao dịch."],
        ["VaR",       "Value at Risk",                        "Đo lường rủi ro về giá trị của danh mục."],
        ["RPS",       "Requests Per Second",                  "Số yêu cầu xử lý mỗi giây."],
        ["WS",        "WebSocket",                            "Giao thức truyền hai chiều full-duplex trên TCP, RFC 6455."],
        ["ES",        "Elasticsearch",                        "Hệ thống tìm kiếm và phân tích phân tán."],
        ["SPA",       "Single Page Application",              "Ứng dụng web trang đơn."],
        ["CEX",       "Centralized Exchange",                 "Sàn giao dịch tập trung."],
        ["DEX",       "Decentralized Exchange",               "Sàn giao dịch phi tập trung."],
        ["Order Book","Sổ lệnh",                              "Cấu trúc lưu các lệnh mua/bán còn mở của một cặp giao dịch."],
        ["Spot",      "Giao dịch giao ngay",                  "Giao dịch mua bán thực tại giá hiện hành."],
        ["Futures",   "Hợp đồng tương lai",                   "Hợp đồng phái sinh, đặc biệt Perpetual Futures không có kỳ hạn."],
        ["Leverage",  "Đòn bẩy",                              "Hệ số nhân vốn, cho phép giao dịch lớn hơn ký quỹ."],
        ["Liquidation","Thanh lý",                             "Đóng vị thế cưỡng chế khi margin không đủ bù lỗ."],
        ["Margin",    "Ký quỹ",                               "Số tiền đặt cọc khi mở vị thế phái sinh."],
        ["Mark Price","Giá đánh dấu",                         "Giá tham chiếu dùng để tính PnL, tránh thao túng."],
        ["Hot path",  "Đường xử lý nóng",                     "Đường code xử lý yêu cầu có độ trễ thấp, khối lượng cao."],
        ["Fan-out",   "Kỹ thuật phát tán",                    "Một sự kiện được phân phối đến nhiều consumer đồng thời."],
        ["Outbox",    "Mẫu thiết kế Outbox",                  "Ghi sự kiện vào bảng outbox cùng transaction để đảm bảo nhất quán."],
        ["Projector", "Bộ chiếu dữ liệu",                    "Consumer tiêu thụ sự kiện và ghi read model vào CSDL."],
    ],
    col_widths=[3.5, 4.5, 8.0],
)
CAP(doc, "Bảng A.1. Danh mục từ viết tắt và thuật ngữ chuyên ngành")
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# PHU LUC B. MAU JSON
# ═══════════════════════════════════════════════════════════════════
H(doc, "PHỤ LỤC B. MẪU JSON REQUEST / RESPONSE", 1)

P(doc, "B.1. Đặt lệnh Spot — Request", bold=True)
P(doc,
    'POST /api/trading/orders\n'
    'Authorization: Bearer <jwt>\n'
    'Content-Type: application/json\n'
    '{\n'
    '  "pair": "BTC_USDT",\n'
    '  "side": "BUY",\n'
    '  "type": "LIMIT",\n'
    '  "price": 65100.50,\n'
    '  "amount": 0.5\n'
    '}',
    indent=False, size=11)

P(doc, "B.2. Đặt lệnh Spot — Response (thành công)", bold=True)
P(doc,
    '{\n'
    '  "success": true,\n'
    '  "data": {\n'
    '    "id": 1234,\n'
    '    "pair": "BTC_USDT",\n'
    '    "side": "BUY",\n'
    '    "type": "LIMIT",\n'
    '    "price": 65100.50,\n'
    '    "amount": 0.5,\n'
    '    "filledAmount": 0.3,\n'
    '    "status": "PARTIAL",\n'
    '    "createdAt": "2026-04-12T10:30:00Z"\n'
    '  }\n'
    '}',
    indent=False, size=11)

P(doc, "B.3. Mở vị thế Futures — Request", bold=True)
P(doc,
    'POST /api/futures/order\n'
    'Authorization: Bearer <jwt>\n'
    '{\n'
    '  "pair": "BTC_USDT",\n'
    '  "side": "LONG",\n'
    '  "leverage": 10,\n'
    '  "size": 0.1,\n'
    '  "takeProfit": 70000,\n'
    '  "stopLoss": 60000\n'
    '}',
    indent=False, size=11)

P(doc, "B.4. Sự kiện trade.executed (Redis Stream)", bold=True)
P(doc,
    '{\n'
    '  "pair": "BTC_USDT",\n'
    '  "buyOrderId": 1234,\n'
    '  "sellOrderId": 5678,\n'
    '  "buyerId": 42,\n'
    '  "sellerId": 88,\n'
    '  "price": 65100.50,\n'
    '  "amount": 0.3,\n'
    '  "total": 19530.15,\n'
    '  "buyerFee": 19.53,\n'
    '  "sellerFee": 19.53,\n'
    '  "side": "BUY"\n'
    '}',
    indent=False, size=11)

P(doc, "B.5. WebSocket message (trades@BTC_USDT)", bold=True)
P(doc,
    '{\n'
    '  "channel": "trades@BTC_USDT",\n'
    '  "data": {\n'
    '    "price": 65100.50,\n'
    '    "amount": 0.3,\n'
    '    "side": "BUY"\n'
    '  }\n'
    '}',
    indent=False, size=11)
PB(doc)


# ═══════════════════════════════════════════════════════════════════
# TAI LIEU THAM KHAO
# ═══════════════════════════════════════════════════════════════════
H(doc, "TÀI LIỆU THAM KHẢO", 1)
refs = [
    "[1]  Newman, S. (2021). Building Microservices: Designing Fine-Grained Systems, 2nd Edition. O'Reilly Media.",
    "[2]  Fowler, M. CQRS — Command Query Responsibility Segregation. https://martinfowler.com/bliki/CQRS.html",
    "[3]  Kleppmann, M. (2017). Designing Data-Intensive Applications. O'Reilly Media.",
    "[4]  Richardson, C. (2018). Microservices Patterns: With Examples in Java. Manning Publications.",
    "[5]  Redis Labs. Redis Streams Documentation. https://redis.io/docs/latest/develop/data-types/streams/",
    "[6]  Apache Software Foundation. Apache Kafka Documentation. https://kafka.apache.org/documentation/",
    "[7]  gRPC Authors. gRPC Documentation — Introduction to gRPC. https://grpc.io/docs/what-is-grpc/introduction/",
    "[8]  The Go Team. The Go Programming Language Specification. https://go.dev/ref/spec",
    "[9]  IETF. RFC 6238 — TOTP: Time-Based One-Time Password Algorithm. 2011.",
    "[10] IETF. RFC 7519 — JSON Web Token (JWT). 2015.",
    "[11] IETF. RFC 6455 — The WebSocket Protocol. 2011.",
    "[12] PostgreSQL Global Development Group. PostgreSQL 16 Documentation. https://www.postgresql.org/docs/16/",
    "[13] Elastic N.V. Elasticsearch 8.x Reference. https://www.elastic.co/guide/en/elasticsearch/reference/8.17/",
    "[14] OWASP Foundation. OWASP Top 10 — 2021. https://owasp.org/Top10/",
    "[15] Binance. Binance API Documentation. https://binance-docs.github.io/apidocs/",
    "[16] Bybit. Bybit V5 API Documentation. https://bybit-exchange.github.io/docs/",
    "[17] Docker, Inc. Docker Compose Documentation. https://docs.docker.com/compose/",
    "[18] CoinGecko. CoinGecko API V3 Documentation. https://www.coingecko.com/en/api/documentation",
    "[19] Gorilla WebSocket. gorilla/websocket — A fast, well-tested Go WebSocket implementation. https://github.com/gorilla/websocket",
    "[20] GORM. The fantastic ORM library for Golang. https://gorm.io/docs/",
]
for r in refs:
    P(doc, r, indent=False, after=4, size=12)


# ═══════════════════════════════════════════════════════════════════
# LUU FILE
# ═══════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"File size: {os.path.getsize(OUT) / 1024:.0f} KB")
